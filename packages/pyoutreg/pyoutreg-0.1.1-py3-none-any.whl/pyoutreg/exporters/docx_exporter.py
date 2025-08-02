"""
Word (docx) exporter for regression results using python-docx.
Creates professional publication-quality tables.
"""

import pandas as pd
import numpy as np
from pathlib import Path
from typing import Optional, List, Any
import warnings

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from ..core.options import OutregOptions


class DocxExporter:
    """Export formatted tables to Word documents with professional styling."""
    
    def __init__(self, options: OutregOptions):
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word export")
        
        self.options = options
    
    def export(self, df: pd.DataFrame, filename: str):
        """Export DataFrame to Word document."""
        
        filepath = Path(filename)
        
        # Handle replace/append logic
        if self.options.append and filepath.exists():
            doc = Document(str(filepath))
            # Add some space before new table
            doc.add_paragraph("")
        else:
            doc = Document()
            # Set default font
            self._set_document_style(doc)
        
        # Add title if specified
        if self.options.title:
            title = doc.add_heading(self.options.title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create and populate table
        table = self._create_table(doc, df)
        
        # Apply styling
        self._apply_table_styling(table, df)
        
        # Add notes
        if not self.options.nonotes:
            self._add_notes(doc)
        
        # Save document
        doc.save(str(filepath))
    
    def _set_document_style(self, doc):
        """Set default document styling."""
        
        style = doc.styles['Normal']
        font = style.font
        font.name = self.options.font_name
        font.size = Pt(self.options.font_size)
        
        # Set margins if landscape
        if self.options.landscape:
            sections = doc.sections
            for section in sections:
                section.orientation = 1  # Landscape
                # Swap width and height
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height
    
    def _create_table(self, doc, df: pd.DataFrame):
        """Create table and populate with data."""
        
        # Create table with appropriate dimensions
        num_rows = len(df) + 1  # +1 for header
        num_cols = len(df.columns)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add header
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            header_cells[i].text = str(col_name)
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data
        for row_idx, (_, row) in enumerate(df.iterrows()):
            table_row = table.rows[row_idx + 1]  # +1 because of header
            for col_idx, value in enumerate(row):
                cell = table_row.cells[col_idx]
                cell.text = str(value) if pd.notna(value) and value != '' else ''
                
                # Set alignment
                if col_idx == 0:  # First column (variable names)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:  # Data columns
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return table
    
    def _apply_table_styling(self, table, df: pd.DataFrame):
        """Apply professional styling to the table."""
        
        # Set table style
        table.style = 'Table Grid'
        
        # Identify different row types for styling
        header_row_idx = 0
        title_rows = []
        stats_rows = []
        note_rows = []
        
        for i, (idx, row) in enumerate(df.iterrows()):
            first_cell = str(row.iloc[0]).strip()
            
            # Identify statistics rows
            if first_cell in ['Observations', 'R-squared', 'F-statistic'] or \
               (self.options.addstat and first_cell in self.options.addstat.keys()):
                stats_rows.append(i + 1)  # +1 because of header
            
            # Identify note rows
            elif 'p<' in first_cell or (self.options.addnote and first_cell == self.options.addnote):
                note_rows.append(i + 1)
        
        # Apply borders and formatting
        for row_idx, row in enumerate(table.rows):
            
            # Header row
            if row_idx == header_row_idx:
                self._add_bottom_border(row)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            # Statistics rows (add top border to first stats row)
            elif row_idx in stats_rows and row_idx == min(stats_rows):
                self._add_top_border(row)
            
            # Adjust font size for note rows
            elif row_idx in note_rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(self.options.font_size - 1)
                            run.italic = True
        
        # Set column widths
        self._set_column_widths(table, df)
    
    def _add_top_border(self, row):
        """Add top border to table row."""
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '4')
            top_border.set(qn('w:space'), '0')
            top_border.set(qn('w:color'), '000000')
            tcBorders.append(top_border)
            tcPr.append(tcBorders)
    
    def _add_bottom_border(self, row):
        """Add bottom border to table row."""
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), '000000')
            tcBorders.append(bottom_border)
            tcPr.append(tcBorders)
    
    def _set_column_widths(self, table, df: pd.DataFrame):
        """Set appropriate column widths."""
        
        # Calculate relative widths
        col_widths = []
        for col_idx, column in enumerate(df.columns):
            if col_idx == 0:  # Variable names column
                col_widths.append(2.5)
            else:  # Data columns
                col_widths.append(1.2)
        
        # Apply widths
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
    
    def _add_notes(self, doc):
        """Add notes to the document."""
        
        # Add space
        doc.add_paragraph("")
        
        # Add custom note if specified
        if self.options.addnote:
            note_para = doc.add_paragraph(self.options.addnote)
            note_para.style.font.size = Pt(self.options.font_size - 1)
            note_para.style.font.italic = True
        
        # Add significance note
        sig_note = doc.add_paragraph("*** p<0.01, ** p<0.05, * p<0.1")
        sig_note.style.font.size = Pt(self.options.font_size - 1)
        sig_note.style.font.italic = True


def export_to_docx(df: pd.DataFrame, filename: str, options: OutregOptions):
    """Convenience function for Word export."""
    exporter = DocxExporter(options)
    exporter.export(df, filename)
