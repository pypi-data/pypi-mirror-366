"""Document readers for various file formats."""

import chardet
from pathlib import Path
from typing import Dict, List, Optional, Type

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from vector_db_query.document_processor.base import DocumentReader
from vector_db_query.document_processor.exceptions import (
    DocumentReadError,
    UnsupportedFileTypeError
)
from vector_db_query.document_processor.excel_reader import ExcelReader
from vector_db_query.document_processor.powerpoint_reader import PowerPointReader
from vector_db_query.document_processor.email_reader import EmailReader
from vector_db_query.document_processor.html_reader import HTMLReader
from vector_db_query.document_processor.config_reader import (
    JSONReader, XMLReader, YAMLReader, INIReader, LogReader
)
from vector_db_query.utils.logger import get_logger

# Import OCR reader conditionally
try:
    from vector_db_query.document_processor.image_ocr_reader import ImageOCRReader
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

logger = get_logger(__name__)


class TextReader(DocumentReader):
    """Reader for plain text and markdown files."""
    
    def can_read(self, file_path: Path) -> bool:
        """Check if this reader can handle the file type."""
        return file_path.suffix.lower() in self.supported_extensions
        
    def read(self, file_path: Path) -> str:
        """Read text file with encoding detection."""
        try:
            # First try to detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # Detect encoding
            detection = chardet.detect(raw_data)
            encoding = detection.get('encoding', 'utf-8')
            confidence = detection.get('confidence', 0)
            
            logger.debug(
                f"Detected encoding for {file_path.name}: "
                f"{encoding} (confidence: {confidence:.2f})"
            )
            
            # Decode with detected encoding
            try:
                text = raw_data.decode(encoding)
            except (UnicodeDecodeError, TypeError):
                # Fall back to utf-8 with error handling
                logger.warning(
                    f"Failed to decode with {encoding}, "
                    f"falling back to utf-8 with replacement"
                )
                text = raw_data.decode('utf-8', errors='replace')
                
            return text.strip()
            
        except Exception as e:
            raise DocumentReadError(
                f"Failed to read text file: {e}",
                file_path=str(file_path)
            )
            
    @property
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['.txt', '.md', '.text', '.markdown']


class PDFReader(DocumentReader):
    """Reader for PDF files."""
    
    def can_read(self, file_path: Path) -> bool:
        """Check if this reader can handle the file type."""
        return file_path.suffix.lower() in self.supported_extensions
        
    def read(self, file_path: Path) -> str:
        """Read PDF file and extract text."""
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(
                        f"Failed to extract text from page {page_num} "
                        f"of {file_path.name}: {e}"
                    )
                    
            if not text_parts:
                raise DocumentReadError(
                    "No text could be extracted from PDF",
                    file_path=str(file_path)
                )
                
            return '\n\n'.join(text_parts).strip()
            
        except Exception as e:
            raise DocumentReadError(
                f"Failed to read PDF file: {e}",
                file_path=str(file_path)
            )
            
    @property
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['.pdf']


class DocxReader(DocumentReader):
    """Reader for Word documents."""
    
    def can_read(self, file_path: Path) -> bool:
        """Check if this reader can handle the file type."""
        return file_path.suffix.lower() in self.supported_extensions
        
    def read(self, file_path: Path) -> str:
        """Read Word document and extract text."""
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
                        
            if not text_parts:
                raise DocumentReadError(
                    "No text could be extracted from document",
                    file_path=str(file_path)
                )
                
            return '\n\n'.join(text_parts).strip()
            
        except Exception as e:
            raise DocumentReadError(
                f"Failed to read Word document: {e}",
                file_path=str(file_path)
            )
            
    @property
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['.docx', '.doc']


class ReaderFactory:
    """Factory for creating appropriate document readers."""
    
    def __init__(self, enable_ocr: bool = False, ocr_language: str = "eng"):
        """Initialize the reader factory.
        
        Args:
            enable_ocr: Whether to enable OCR for image files
            ocr_language: OCR language code (default: 'eng')
        """
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language
        
        self._readers: Dict[str, DocumentReader] = {
            'text': TextReader(),
            'pdf': PDFReader(),
            'docx': DocxReader(),
            'excel': ExcelReader(),
            'powerpoint': PowerPointReader(),
            'email': EmailReader(),
            'html': HTMLReader(),
            'json': JSONReader(),
            'xml': XMLReader(),
            'yaml': YAMLReader(),
            'ini': INIReader(),
            'log': LogReader()
        }
        
        # Add OCR reader if available and enabled
        if HAS_OCR and enable_ocr:
            self._readers['ocr'] = ImageOCRReader(language=ocr_language)
            logger.info(f"OCR support enabled with language: {ocr_language}")
        elif HAS_OCR:
            logger.info("OCR available but not enabled")
        else:
            logger.warning("OCR support not available - install pillow and pytesseract")
        
        # Build extension to reader mapping
        self._extension_map: Dict[str, DocumentReader] = {}
        for reader in self._readers.values():
            for ext in reader.supported_extensions:
                self._extension_map[ext.lower()] = reader
                
    def get_reader(self, file_path: Path) -> DocumentReader:
        """Get appropriate reader for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Appropriate document reader
            
        Raises:
            UnsupportedFileTypeError: If no reader supports the file type
        """
        extension = file_path.suffix.lower()
        reader = self._extension_map.get(extension)
        
        if reader is None:
            raise UnsupportedFileTypeError(
                f"No reader available for file type: {extension}",
                file_path=str(file_path)
            )
            
        return reader
        
    def read_document(self, file_path: Path, enable_ocr: Optional[bool] = None, ocr_language: Optional[str] = None) -> str:
        """Read a document using the appropriate reader.
        
        Args:
            file_path: Path to the file
            enable_ocr: Whether to enable OCR (overrides instance setting)
            ocr_language: OCR language code (overrides instance setting)
            
        Returns:
            Text content of the document
            
        Raises:
            DocumentReadError: If document cannot be read
            UnsupportedFileTypeError: If file type is not supported
        """
        # Use provided values or fall back to instance settings
        use_ocr = enable_ocr if enable_ocr is not None else self.enable_ocr
        ocr_lang = ocr_language if ocr_language is not None else self.ocr_language
        
        # For image files, check if we should use OCR
        extension = file_path.suffix.lower()
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif'}
        
        if extension in image_extensions and use_ocr and HAS_OCR:
            # Use OCR reader for images when OCR is enabled
            if 'ocr' not in self._readers:
                # Create OCR reader on demand if not already created
                from vector_db_query.document_processor.image_ocr_reader import ImageOCRReader
                self._readers['ocr'] = ImageOCRReader(language=ocr_lang)
            reader = self._readers['ocr']
            logger.info(f"Reading {file_path.name} with OCR (language: {ocr_lang})")
        else:
            # Use standard reader
            reader = self.get_reader(file_path)
            logger.info(f"Reading {file_path.name} with {reader.__class__.__name__}")
            
        return reader.read(file_path)
        
    @property
    def supported_extensions(self) -> List[str]:
        """Get all supported file extensions."""
        extensions = list(self._extension_map.keys())
        # Add image extensions if OCR is available
        if HAS_OCR or self.enable_ocr:
            image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']
            for ext in image_extensions:
                if ext not in extensions:
                    extensions.append(ext)
        return extensions
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get all potentially supported file extensions.
        
        Returns:
            List of supported extensions (without leading dots)
        """
        # Create a temporary instance to get all extensions
        factory = cls(enable_ocr=True)
        return [ext.lstrip('.') for ext in factory.supported_extensions]
        
    def register_reader(self, name: str, reader: DocumentReader) -> None:
        """Register a custom reader.
        
        Args:
            name: Name for the reader
            reader: DocumentReader instance
        """
        self._readers[name] = reader
        for ext in reader.supported_extensions:
            self._extension_map[ext.lower()] = reader
            
        logger.info(
            f"Registered reader '{name}' for extensions: "
            f"{reader.supported_extensions}"
        )