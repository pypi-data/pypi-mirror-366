#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SAST漏洞修复服务
提供docx文档的SAST报告解析、漏洞修复建议生成和报告导出功能
"""
import csv
import os
import tempfile
import logging
import asyncio
import json
import re
from contextlib import asynccontextmanager
from typing import AsyncIterator, Dict, Any, Optional, List

from mcp.server.models import InitializationOptions
import mcp.types as types
from mcp.server import NotificationOptions, Server
import mcp.server.stdio
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "sast_fixer_mcp.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("SASTFixerMCPServer")

# 创建MCP服务器
server = Server("sast-fixer-mcp")

def get_paragraph_heading_level(paragraph):
    """获取段落标题级别"""
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.replace('Heading', ''))
        except ValueError:
            return 0
    return 0

def process_table(table):
    """处理表格数据"""
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_text += paragraph.text + " "
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def get_docx_elements_in_order(document):
    """按顺序获取文档元素"""
    body = document._element.body
    elements = []
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            for paragraph in document.paragraphs:
                if paragraph._element is child:
                    elements.append(('paragraph', paragraph))
                    break
        elif child.tag.endswith('tbl'):
            for table in document.tables:
                if table._element is child:
                    elements.append(('table', table))
                    break
    return elements

def parse_docx_to_json(docx_path):
    """解析docx文档为JSON结构"""
    doc = Document(docx_path)
    elements = get_docx_elements_in_order(doc)
    result = []
    heading_stack = [(-1, None, result)]
    current_content = []

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            level = get_paragraph_heading_level(elem)
            
            if level > 0:
                if heading_stack[-1][1] is not None:
                    heading_stack[-1][1]["content"] = current_content
                
                heading_node = {
                    "title": elem.text,
                    "content": []
                }
                current_content = []
                
                while heading_stack[-1][0] >= level:
                    heading_stack.pop()
                
                parent_level, parent_node, parent_container = heading_stack[-1]
                
                if isinstance(parent_container, list):
                    parent_container.append(heading_node)
                else:
                    if "children" not in parent_container:
                        parent_container["children"] = []
                    parent_container["children"].append(heading_node)
                
                heading_stack.append((level, heading_node, heading_node))
            elif elem.text.strip():
                current_content.append({"type": "text", "value": elem.text})
        
        elif elem_type == 'table':
            table_data = process_table(elem)
            current_content.append({"type": "table", "value": table_data})
    
    if heading_stack[-1][1] is not None:
        heading_stack[-1][1]["content"] = current_content
    
    return result

def transform_json(json_data):
    """转换JSON结构为漏洞报告格式"""
    result = {}
    issue_level_pattern = re.compile(r"【(低危|中危|高危|提示)】.*?漏洞数：(\d+)")

    for item in json_data:
        main_title = item["title"]
        result[main_title] = []
        if "children" not in item:
            continue
            
        for child in item["children"]:
            match = issue_level_pattern.search(child["title"])
            if match:
                issue_level = match.group(1)
                issue_count = match.group(2)

                if issue_level == "提示":
                    issue_level = "Notice"
                elif issue_level == "低危":
                    issue_level = "Low"
                elif issue_level == "中危":
                    issue_level = "Medium"
                elif issue_level == "高危":
                    issue_level = "High"
                    
                if issue_level not in ["Medium", "High"]:
                    continue

                issue = {
                    "issue_title": child["title"],
                    "issue_level": issue_level,
                    "issue_count": issue_count,
                    "issue_desc": "",
                    "fix_advice": "",
                    "code_sample": "",
                    "code_list": []
                }
                
                for section in child["children"]:
                    section_title = section["title"]
                    section_content = section["content"]
                    
                    if section_title == "漏洞描述":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["issue_desc"] = content_item["value"]
                    
                    elif section_title == "修复建议":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["fix_advice"] = content_item["value"]
                    
                    elif section_title == "代码示例":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["code_sample"] = content_item["value"]
                    
                    elif re.match(r'^NO\.\d+\.\s代码位置$', section_title):
                        code_item = {}
                        
                        if section_content and section_content[0]["type"] == "text":
                            code_location_num = section_content[0]["value"]
                            splitor = code_location_num.split(":")
                            code_item["code_location"] = splitor[0]
                            code_item["code_line_num"] = splitor[1]
                        
                        if "children" in section:
                            for child_section in section["children"]:
                                for content_item in child_section["content"]:
                                    if content_item["type"] == "table":
                                        code_item["code_details"] = content_item["value"][-1][1]
                        
                        issue["code_list"].append(code_item)
                
                result[main_title].append(issue)
    
    return result

def save_json(data, output_path):
    """保存JSON数据到文件"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON数据已保存到{output_path}")

@server.list_tools()
async def handle_list_tools() -> List[types.Tool]:
    """
    列出可用的SAST报告处理工具
    """
    return [
        types.Tool(
            name="convert_sast_docx_to_json",
            description="将SAST报告的docx文档转换为JSON格式",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "SAST报告docx文件路径"
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="get_pending_vulnerability_json_files",
            description="获取.scanissuefix目录中所有待处理的漏洞JSON文件(_new.json)",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        types.Tool(
            name="generate_csv_report",
            description="从所有已完成的漏洞JSON文件(_finished.json)生成CSV报告",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        )
    ]

@server.call_tool()
async def handle_call_tool(
    name: str, arguments: Dict[str, Any]
) -> List[types.TextContent]:
    """
    处理工具执行请求
    """
    try:
        if name == "convert_sast_docx_to_json":
            file_path = arguments["file_path"]
            if not os.path.exists(file_path):
                return [types.TextContent(type="text", text=f"文件不存在: {file_path}")]
            
            # 解析docx为JSON
            all_docx_json = parse_docx_to_json(file_path)
            result = []
            for item in all_docx_json:
                if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]: 
                    result += transform_json([item])[item["title"]]

            # Prepare the final result
            json_result = {"status": "success", "data": result}

            # 确保输出目录存在
            output_dir = ".scanissuefix"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # Iterate over each issue in the 'data' section
            for index, issue in enumerate(json_result["data"], start=1):
                issue_title = issue["issue_title"]

                # Check if the code list has more than 5 entries
                code_list = issue["code_list"]
                total_code_entries = len(code_list)

                if total_code_entries > 5:
                    # Split the code list into chunks of 5 or fewer
                    chunk_size = 5
                    chunks = [code_list[i:i + chunk_size] for i in range(0, total_code_entries, chunk_size)]
                    
                    # Create separate files for each chunk
                    for chunk_index, chunk in enumerate(chunks, start=1):
                        issue_copy = issue.copy()
                        issue_copy["code_list"] = chunk

                        # Construct the filename based on the issue title and index, with chunk numbering
                        filename = f"{index}_{issue_title.replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace('/', '')}_{chunk_index}_new.json"
                        file_path = os.path.join(output_dir, filename)

                        # Save each issue's details to a new file
                        with open(file_path, 'w', encoding='utf-8') as f:
                            json.dump(issue_copy, f, ensure_ascii=False, indent=2)

                        print(f"Issue saved to: {file_path}")
                else:
                    # If there are 5 or fewer code entries, save as a single file
                    filename = f"{index}_{issue_title.replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace('/', '')}_new.json"
                    file_path = os.path.join(output_dir, filename)
                    # Save each issue's details to a new file
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(issue, f, ensure_ascii=False, indent=2)

            
            return [types.TextContent(type="text", text=f"所有漏洞已保存到{output_dir}目录")]

        elif name == "get_pending_vulnerability_json_files":
            output_dir = ".scanissuefix"
            if not os.path.exists(output_dir):
                return [types.TextContent(type="text", text=f"目录'{output_dir}'不存在")]
            
            all_files = os.listdir(output_dir)
            new_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_new.json")]

            if not new_json_files:
                return [types.TextContent(type="text", text=f"目录'{output_dir}'中没有找到'_new.json'文件")]
            
            return [types.TextContent(type="text", text="\n".join(new_json_files))]

        elif name == "generate_csv_report":
            output_dir = ".scanissuefix"
            if not os.path.exists(output_dir):
                return [types.TextContent(type="text", text=f"目录'{output_dir}'不存在")]
            
            all_files = os.listdir(output_dir)
            finished_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_finished.json")]

            if not finished_json_files:
                return [types.TextContent(type="text", text=f"目录'{output_dir}'中没有找到'_finished.json'文件")]
            
            # 准备CSV报告数据
            report_data = []
            for file_path in finished_json_files:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                    for code_item in data.get("code_list", []):
                        report_data.append({
                            "漏洞类型": data["issue_title"],
                            "漏洞等级": data["issue_level"],
                            "代码位置": code_item["code_location"],
                            "代码行号": code_item["code_line_num"],
                            "代码详情": code_item["code_details"],
                            "修复状态": code_item.get("status", "missed"),
                            "误报概率": code_item.get("false_positive_probability", ""),
                            "误报原因澄清": code_item.get("false_positive_reason", "")
                        })

            # 定义CSV输出文件路径
            csv_output_path = os.path.join(output_dir, "sast_fix_report.csv")

            # 写入CSV文件
            with open(csv_output_path, 'w', newline='', encoding='utf-8') as csv_file:
                fieldnames = [
                    "漏洞类型", 
                    "漏洞等级", 
                    "代码位置", 
                    "代码行号", 
                    "代码详情", 
                    "修复状态", 
                    "误报概率", 
                    "误报原因澄清"
                ]
                writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL)
                writer.writeheader()
                writer.writerows(report_data)

            return [types.TextContent(type="text", text=f"CSV报告生成成功: {csv_output_path}")]

        else:
            raise ValueError(f"未知工具: {name}")
    except Exception as e:
        error_msg = f"工具执行失败: {str(e)}"
        logger.error(error_msg)
        return [types.TextContent(type="text", text=error_msg)]

async def main():
    """运行MCP服务器"""
    async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="sast-fixer-mcp",
                server_version="0.1.0",
                capabilities=server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                ),
            ),
        )

if __name__ == "__main__":
    # 启动时删除可能存在的旧状态文件
    current_doc_file = os.path.join(tempfile.gettempdir(), "docx_mcp_current_doc.txt")
    if os.path.exists(current_doc_file):
        try:
            os.remove(current_doc_file)
            logger.info("已移除旧的状态文件")
        except Exception as e:
            logger.error(f"移除旧状态文件失败: {e}")
    
    # 运行MCP服务器
    asyncio.run(main())