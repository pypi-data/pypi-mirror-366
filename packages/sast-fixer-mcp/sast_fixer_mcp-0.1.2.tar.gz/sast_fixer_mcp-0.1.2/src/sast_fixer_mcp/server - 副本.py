#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
MCP Docx Processing Service
Provides various operations for docx documents, including querying, adding, modifying, deleting, and font style settings
Implemented using the official MCP library
"""
import csv
import os
import tempfile
import logging
import traceback
from contextlib import asynccontextmanager
from typing import AsyncIterator, Dict, Any, Optional

from mcp.server.fastmcp import FastMCP, Context
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "docx_mcp_server.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("DocxMCPServer")

# Create a state file for restoring state when MCP service restarts
CURRENT_DOC_FILE = os.path.join(tempfile.gettempdir(), "docx_mcp_current_doc.txt")

class DocxProcessor:
    """Class for processing Docx documents, implementing various document operations"""
    
    def __init__(self):
        self.documents = {}  # Store opened documents
        self.current_document = None
        self.current_file_path = None
        
        # Try to load current document from state file
        self._load_current_document()
    
    def _load_current_document(self):
        """Load current document from state file"""
        if not os.path.exists(CURRENT_DOC_FILE):
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'r', encoding='utf-8') as f:
                file_path = f.read().strip()
            
            if file_path and os.path.exists(file_path):
                try:
                    self.current_file_path = file_path
                    self.current_document = Document(file_path)
                    self.documents[file_path] = self.current_document
                    return True
                except Exception as e:
                    logger.error(f"Failed to load document at {file_path}: {e}")
                    # Delete invalid state file to prevent future loading attempts
                    try:
                        os.remove(CURRENT_DOC_FILE)
                        logger.info(f"Removed invalid state file pointing to {file_path}")
                    except Exception as e_remove:
                        logger.error(f"Failed to remove state file: {e_remove}")
            else:
                # Delete invalid state file if path is empty or file doesn't exist
                try:
                    os.remove(CURRENT_DOC_FILE)
                    logger.info("Removed invalid state file with non-existent document path")
                except Exception as e_remove:
                    logger.error(f"Failed to remove state file: {e_remove}")
        except Exception as e:
            logger.error(f"Failed to load current document: {e}")
            # Delete corrupted state file
            try:
                os.remove(CURRENT_DOC_FILE)
                logger.info("Removed corrupted state file")
            except Exception as e_remove:
                logger.error(f"Failed to remove state file: {e_remove}")
        
        return False
    
    def _save_current_document(self):
        """Save current document path to state file"""
        if not self.current_file_path:
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'w', encoding='utf-8') as f:
                f.write(self.current_file_path)
            return True
        except Exception as e:
            logger.error(f"Failed to save current document path: {e}")
        
        return False
    
    def save_state(self):
        """Save processor state"""
        # Save current document
        if self.current_document and self.current_file_path:
            try:
                self.current_document.save(self.current_file_path)
                self._save_current_document()
            except Exception as e:
                logger.error(f"Failed to save current document: {e}")
    
    def load_state(self):
        """Load processor state"""
        self._load_current_document()

    # ... Keep all original document processing methods ...

# Create global processor instance
processor = DocxProcessor()

@asynccontextmanager
async def server_lifespan(server: FastMCP) -> AsyncIterator[Dict[str, Any]]:
    """Manage server lifecycle"""
    try:
        # Start server with clean state
        logger.info("DocxProcessor MCP server starting with clean state...")
        # Do not attempt to load any previous state
        yield {"processor": processor}
    finally:
        # Save state when server shuts down
        logger.info("DocxProcessor MCP server shutting down...")
        if processor.current_document and processor.current_file_path:
            processor.save_state()
        else:
            logger.info("No document open, not saving state")

# Create MCP server
mcp = FastMCP(
    name="DocxProcessor",
    instructions="Word document processing service, providing functions to create, edit, and query documents",
    lifespan=server_lifespan
)

@mcp.tool()
def create_document(ctx: Context, file_path: str) -> str:
    """
    Create a new Word document
    
    Parameters:
    - file_path: Document save path
    """
    try:
        processor.current_document = Document()
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        # Save document
        processor.current_document.save(file_path)
        
        return f"Document created successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to create document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def open_document(ctx: Context, file_path: str) -> str:
    """
    Open an existing Word document
    
    Parameters:
    - file_path: Path to the document to open
    """
    try:
        if not os.path.exists(file_path):
            return f"File does not exist: {file_path}"
        
        processor.current_document = Document(file_path)
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        return f"Document opened successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to open document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def save_document(ctx: Context) -> str:
    """
    Save the currently open Word document to the original file (update the original file)
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        if not processor.current_file_path:
            return "Current document has not been saved before, please use save_as_document to specify a save path"
            
        # Save to original file path
        processor.current_document.save(processor.current_file_path)
        
        return f"Document saved successfully to original file: {processor.current_file_path}"
    except Exception as e:
        error_msg = f"Failed to save document: {str(e)}"
        logger.error(error_msg)
        return error_msg




# Add more tools...


import json
from docx import Document
import re
import os
from pydantic import BaseModel
from typing import Dict

def get_paragraph_heading_level(paragraph):
    
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.replace('Heading', ''))
        except ValueError:
            return 0
    return 0

def process_table(table):
    
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:

            cell_text = ""
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_text += paragraph.text + " "
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def get_docx_elements_in_order(document):

    body = document._element.body
    
    elements = []
    for child in body.iterchildren():
        if child.tag.endswith('p'): 
            for paragraph in document.paragraphs:
                if paragraph._element is child:
                    elements.append(('paragraph', paragraph))
                    break
        elif child.tag.endswith('tbl'): 
            for table in document.tables:
                if table._element is child:
                    elements.append(('table', table))
                    break
    
    return elements

def parse_docx_to_json(docx_path):
    
    doc = Document(docx_path)
    elements = get_docx_elements_in_order(doc)
    result = []
    heading_stack = [(-1, None, result)]  # (level, heading_node, container)
    current_content = []
    

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            level = get_paragraph_heading_level(elem)
            
            if level > 0:  # 如果是标题

                if heading_stack[-1][1] is not None:
                    heading_stack[-1][1]["content"] = current_content
                

                heading_node = {
                    "title": elem.text,
                    "content": []
                }
                current_content = []
                

                while heading_stack[-1][0] >= level:
                    heading_stack.pop()
                

                parent_level, parent_node, parent_container = heading_stack[-1]
                
                if isinstance(parent_container, list):
                    parent_container.append(heading_node)
                else:
                    if "children" not in parent_container:
                        parent_container["children"] = []
                    parent_container["children"].append(heading_node)
                

                heading_stack.append((level, heading_node, heading_node))
            elif elem.text.strip():  # 普通段落
                current_content.append({"type": "text", "value": elem.text})
        
        elif elem_type == 'table':

            table_data = process_table(elem)
            current_content.append({"type": "table", "value": table_data})
    

    if heading_stack[-1][1] is not None:
        heading_stack[-1][1]["content"] = current_content
    
    return result

def transform_json(json_data):
    result = {}
    
    # Regex to extract the vulnerability level and number of vulnerabilities
    issue_level_pattern = re.compile(r"【(低危|中危|高危|提示)】.*?漏洞数：(\d+)")

    for item in json_data:
        main_title = item["title"]
        result[main_title] = []
        if "children" not in item:
            continue
        for child in item["children"]:
            # Extract risk level and vulnerability count from the title
            match = issue_level_pattern.search(child["title"])
            if match:
                issue_level = match.group(1)
                issue_count = match.group(2)

                # Map risk level to English
                if issue_level == "提示":
                    issue_level = "Notice"
                elif issue_level == "低危":
                    issue_level = "Low"
                elif issue_level == "中危":
                    issue_level = "Medium"
                elif issue_level == "高危":
                    issue_level = "High"
            if issue_level not in ["Medium", "High"]:
                continue

            issue = {
                "issue_title": child["title"],
                "issue_level": issue_level,  # Added risk level
                "issue_count": issue_count,  # Added vulnerability count
                "issue_desc": "",
                "fix_advice": "",
                "code_sample": "",
                "code_list": []
                # "issue_list": []
            }
            
            # Prepare an issue item to hold all the details
            # issue_item = {
            #     "issue_desc": "",
            #     "fix_advice": "",
            #     "code_sample": "",
            #     "code_list": []
            # }
            
            # Extract content from child's children
            for section in child["children"]:
                section_title = section["title"]
                section_content = section["content"]
                
                # Handle regular sections
                if section_title == "漏洞描述":
                    for content_item in section_content:
                        if content_item["type"] == "text":
                            issue["issue_desc"] = content_item["value"]
                
                elif section_title == "修复建议":
                    for content_item in section_content:
                        if content_item["type"] == "text":
                            issue["fix_advice"] = content_item["value"]
                
                elif section_title == "代码示例":
                    for content_item in section_content:
                        if content_item["type"] == "text":
                            issue["code_sample"] = content_item["value"]
                
                # Handle code location sections
                elif re.match(r'^NO\.\d+\.\s代码位置$', section_title):
                    code_item = {}
                    

                    code_location_num = ""
                    # Get code details from children
                    if "children" in section:
                        for child_section in section["children"]:
                            for content_item in child_section["content"]:
                                if content_item["type"] == "table":
                                    # Get the code details from the table
                                    code_item["code_details"] = content_item["value"][len(content_item["value"]) - 1][1]
                                    if len(section_content)<1:
                                        code_location_num = content_item["value"][len(content_item["value"]) - 1][0]

                    # Get the code location
                    if section_content and section_content[0]["type"] == "text" and section_content!=[]:
                        code_location_num = section_content[0]["value"]
                    splitor = code_location_num.split(":")
                    code_item["code_location"] = splitor[0]
                    code_item["code_line_num"] = splitor[1]
                    
                    
                    
                    issue["code_list"].append(code_item)
            
            # Add the issue item to the issues list
            # issue["issue_list"].append(issue_item)
            result[main_title].append(issue)
    
    return result

def save_json(data, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"JSON data saved to {output_path}")


# Define the input parameters as Pydantic model for validation
class IssueData(BaseModel):
    iss_desc: str
    fix_advice: str
    code_sample: str
    code_location: str
    code_line_num: str
    code_details: str
    code_snippet: str



# @mcp.tool()
# def convert_sast_docx_to_json(ctx: Context, file_path:str):
#     """
#     打开Word文档，将docx版本的SAST报告，解析漏洞，生成漏洞列表的json
#     Parameters:
#     - file_path: Path to the document to open
#     """

#     try:
#         all_docx_json = parse_docx_to_json(file_path)
#         result = []
#         for item in all_docx_json:
#             if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]: 
#                 result += transform_json([item])[item["title"]]

#         json_result = {"status": "success", "data": result}
#         # save_json(json_result, "json_result.json")
#         return json_result
    
#     except Exception as e:
#         return {"status": "error", "message": str(e)}
        
        


@mcp.tool()
def convert_sast_docx_to_json(ctx: Context, file_path: str):
    """
    打开Word文档，将docx版本的SAST报告，解析漏洞，生成漏洞列表的json，并将每个漏洞保存成一个文件。
    Parameters:
    - file_path: Path to the document to open
    """
    try:
        # Parse the docx to JSON
        all_docx_json = parse_docx_to_json(file_path)
        result = []
        for item in all_docx_json:
            if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]: 
                result += transform_json([item])[item["title"]]

        # Prepare the final result
        json_result = {"status": "success", "data": result}

        # Ensure the directory exists
        output_dir = ".scanissuefix"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Iterate over each issue in the 'data' section
        for index, issue in enumerate(json_result["data"], start=1):
            issue_title = issue["issue_title"]

            # Check if the code list has more than 5 entries
            code_list = issue["code_list"]
            total_code_entries = len(code_list)

            if total_code_entries > 5:
                # Split the code list into chunks of 5 or fewer
                chunk_size = 5
                chunks = [code_list[i:i + chunk_size] for i in range(0, total_code_entries, chunk_size)]
                
                # Create separate files for each chunk
                for chunk_index, chunk in enumerate(chunks, start=1):
                    issue_copy = issue.copy()
                    issue_copy["code_list"] = chunk

                    # Construct the filename based on the issue title and index, with chunk numbering
                    filename = f"{index}_{issue_title.replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace('/', '')}_{chunk_index}_new.json"
                    file_path = os.path.join(output_dir, filename)

                    # Save each issue's details to a new file
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(issue_copy, f, ensure_ascii=False, indent=2)

                    print(f"Issue saved to: {file_path}")
            else:
                # If there are 5 or fewer code entries, save as a single file
                filename = f"{index}_{issue_title.replace('【', '_').replace('】', '').replace('(', '_').replace(')', '').replace('：', '').replace('/', '')}_new.json"
                file_path = os.path.join(output_dir, filename)

                # Save each issue's details to a new file
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(issue, f, ensure_ascii=False, indent=2)

                print(f"Issue saved to: {file_path}")

        return {"status": "success", "message": f"All issues saved to .scanissuefix directory"}
    
    except Exception as e:
        return {"status": "error", "message": str(e)}

@mcp.tool()
def get_pending_vulnerability_json_files(ctx: Context) -> str:
    """
    Get all file paths in the .scanissuefix directory that end with '_new.json' (indicating pending vulnerabilities).
    """
    try:
        output_dir = ".scanissuefix"
        # Check if the directory exists
        if not os.path.exists(output_dir):
            return f"The directory '{output_dir}' does not exist."
        
        # List all files in the directory
        all_files = os.listdir(output_dir)
        new_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_new.json")]

        if not new_json_files:
            return f"No '_new.json' files found in the directory '{output_dir}'."
        
        # Return all the found paths
        return "\n".join(new_json_files)
    
    except Exception as e:
        error_msg = f"Failed to get file paths: {str(e)}"
        logger.error(error_msg)
        return error_msg




@mcp.tool()
def generate_csv_report(ctx: Context) -> str:
    """
    Generate a final CSV report (sast_fix_report.csv) from all files ending with '_finished.json' in the '.scanissuefix' directory.
    The CSV report will contain the following columns: 
    "漏洞类型", "漏洞等级", "代码位置", "代码行号", "代码详情", "修复状态", "误报概率", "误报原因澄清".
    """
    try:
        # Directory containing the finished JSON files
        output_dir = ".scanissuefix"
        
        # Check if the directory exists
        if not os.path.exists(output_dir):
            return f"The directory '{output_dir}' does not exist."
        
        # List all files in the directory that end with '_finished.json'
        all_files = os.listdir(output_dir)
        finished_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_finished.json")]

        if not finished_json_files:
            return f"No '_finished.json' files found in the directory '{output_dir}'."
        
        # Prepare CSV report data
        report_data = []

        for file_path in finished_json_files:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

                # Extract information for each vulnerability entry in the file
                for code_item in data.get("code_list", []):
                    # Handle missing fields with defaults
                    false_positive_probability = code_item.get("false_positive_probability", "")
                    false_positive_reason = code_item.get("false_positive_reason", "")
                    status = code_item.get("status", "missed")  # Default to "missed" if status is missing

                    report_data.append({
                        "漏洞类型": data["issue_title"],
                        "漏洞等级": data["issue_level"],
                        "代码位置": code_item["code_location"],
                        "代码行号": code_item["code_line_num"],
                        "代码详情": code_item["code_details"],
                        "修复状态": status,
                        "误报概率": false_positive_probability,
                        "误报原因澄清": false_positive_reason
                    })

        # Define the CSV output file path
        csv_output_path = os.path.join(output_dir, "sast_fix_report.csv")

        # Write the data to a CSV file
        with open(csv_output_path, 'w', newline='', encoding='utf-8') as csv_file:
            fieldnames = [
                "漏洞类型", 
                "漏洞等级", 
                "代码位置", 
                "代码行号", 
                "代码详情", 
                "修复状态", 
                "误报概率", 
                "误报原因澄清"
            ]
            writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL)

            # Write the header row
            writer.writeheader()
            # Write the data rows
            writer.writerows(report_data)

        return f"CSV report generated successfully: {csv_output_path}"

    except Exception as e:
        error_msg = f"Failed to generate CSV report: {str(e)}"
        logger.error(error_msg)
        return error_msg







if __name__ == "__main__":
    # Always start with a clean state, don't try to load any previous document
    if os.path.exists(CURRENT_DOC_FILE):
        try:
            os.remove(CURRENT_DOC_FILE)
            logger.info("Removed existing state file for clean startup")
        except Exception as e:
            logger.error(f"Failed to remove existing state file: {e}")
    
    # Run MCP server
    mcp.run() 