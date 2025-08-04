# -*- coding: utf-8 -*-
"""
Created on Sat Jun 19 11:58:32 2021

@author: 1
"""
import docx
from PyPDF2 import PdfFileWriter, PdfFileReader
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from win32com import client
from docx import Document
import numpy as np
import fitz  # pip install PyMuPDF
import string
import os


'''
如果出现 PyPDF2.utils.PdfReadError: Illegal character in Name Object
解决办法:修改PyPDF2模块中的源码，使其能够处理多种编码
步骤1
点击报错日志最上方的generic.py文件，错误提示如下
Traceback (most recent call last):
  File "E:\Program_Practice\Python\pachong\venv\lib\site-packages\PyPDF2\generic.py", line 484, in readFromStream
    return NameObject(name.decode('utf-8'))
UnicodeDecodeError: 'utf-8' codec can't decode byte 0xcb in position 8: invalid continuation byte
步骤2
将如下代码粘贴到步骤1中打开的generic.py文件的486行，然后保存文件

将：
            try:
                ret = name.decode('utf-8')
            except (UnicodeEncodeError, UnicodeDecodeError) as e:
                ret = name.decode('gbk')
            return NameObject(ret)

修改为：

            try:
                ret = name.decode('utf-8')
            except (UnicodeEncodeError, UnicodeDecodeError) as e:
                try:
                    ret name.decode("utf-8")
                except (UnicodeEncodeError, UnicodeDecodeError) as e:
                    ret = name.decode('gbk')
                return NameObject(ret)


步骤3：
此时执行我们自己写的代码后，错误提示变为：
File "E:\Program_Practice\Python\pachong\venv\lib\site-packages\PyPDF2\utils.py", line 238, in b_
    r = s.encode('latin-1')
UnicodeEncodeError: 'latin-1' codec can't encode characters in position 8-9: ordinal not in range(256)

步骤4
点击报错日志中utils.py文件进行修改
修改前：
            try:
                r = s.encode('latin-1')
            if len(s) < 2:
                bc[s] = r
修改后：
            try:
                r = s.encode('latin-1')
            except Exception as e:
                r = s.encode('utf-8')
            if len(s) < 2:
                bc[s] = r

'''



doc = docx.Document()


def pdf_withmarks(data_path, doc_path, pdf_path, pdf_marks_path):
    doc = docx.Document()

    # 添加标题
    def AddHeadText(text, size, level):
        title_ = doc.add_heading('', level)
        # title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
        title_run = title_.add_run(text)  # 添加标题内容
        title_run.font.size = Pt(size)  # 设置标题字体大小
        title_run.font.name = '宋体'  # 设置标题西文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置标题中文字体
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色

    def generateTitleCode(level):
        til = ''
        while level > 0:
            if len(til) < 1:
                til = str(level)
                print(til)
            else:
                til = str(level) + '.' + til
                print(til)
            level -= 1
        return til

    # folderList_1 = os.listdir(path_root+'/'+folderName_0)

    # tilCodeMap={"1":0}
    # 生成文档内容
    def traverse(f, level, tilcode):
        tilCodeMap = {"1": 0}
        fs = os.listdir(f)
        contents_0 = []
        for f1 in fs:
            tmp_path = os.path.join(f, f1)
            if not os.path.isdir(tmp_path):
                if f1.endswith('.png'):
                    pic = doc.add_picture(tmp_path, width=Inches(5))  # 添加图片
                if f1.endswith('.txt'):
                    with open(tmp_path, encoding='utf-8', errors='ignore') as content_1:
                        data_1 = content_1.read()
                        p = doc.add_paragraph(data_1)
                        p.paragraph_format.first_line_indent = Cm(0.8)
                        p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
                        if data_1.startswith('图'):
                            # p = doc.add_paragraph(data_1)
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if tilcode not in tilCodeMap.keys():
                    tilCodeMap[tilcode] = 0
                tilCodeMap[tilcode] = tilCodeMap[tilcode] + 1
                if tilcode == '1':
                    AddHeadText(text=str(tilCodeMap[tilcode]) + '.' + f1[2:], size=12, level=level)
                    # print(str(tilCodeMap[tilcode]) + '.' + f1[2:])
                    contents_1 = str(tilCodeMap[tilcode]) + '.' + f1[2:]
                    contents_0.append(contents_1)
                    # print(contents_0)
                    traverse(tmp_path, level + 1, str(tilCodeMap[tilcode]) + '.')
                else:
                    AddHeadText(text=tilcode + str(tilCodeMap[tilcode]) + '.' + f1[2:], size=12, level=level)
                    # print(tilcode + str(tilCodeMap[tilcode]) + '.' + f1[2:])
                    contents_2 = tilcode + str(tilCodeMap[tilcode]) + '.' + f1[2:]
                    contents_0.append(contents_2)
                    # print(contents_0)
                    traverse(tmp_path, level + 1, tilcode + str(tilCodeMap[tilcode]) + '.')

    # word转pdf
    def doc2pdf(doc_name, pdf_name):

        try:
            word = client.DispatchEx("Word.Application")
            if os.path.exists(pdf_name):
                os.remove(pdf_name)
            worddoc = word.Documents.Open(doc_name, ReadOnly=1)
            worddoc.SaveAs(pdf_name, FileFormat=17)
            worddoc.Close()
            return pdf_name
        except:
            return 0

            # path_root = r'G:/Work/MODE/公报/公报转pdf测试数据'

    path_root = data_path
    # 正文标题
    folderList_0 = os.listdir(path_root)
    folderName_0 = folderList_0[0]
    title_0 = doc.add_paragraph()  # 正文大标题
    run = title_0.add_run(folderName_0)  # 使用add_run添加文字
    run.font.size = Pt(20)  # 设置字体大小
    title_0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    run.bold = True  # 加粗
    doc.styles['Normal'].font.name = '宋体'  # 设置字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    path = path_root + '/' + folderName_0
    # 调用traverse函数
    traverse(path, 1, "1")
    doc.save(doc_path)
    output_pdf = pdf_path

    # 调用doc2pdf函数,将doc转为pdf
    rc = doc2pdf(doc_path, output_pdf)

    # 生成目录列表catalog
    catalog = []
    for ii in np.arange(10):
        # print(ii)
        document = Document(doc_path)
        all_paragraphs = doc.paragraphs
        for paragraph in all_paragraphs:
            # 打印每一个段落的文字
            # print(paragraph.text)
            if paragraph.text.startswith(str(ii) + '.'):
                catalog_0 = paragraph.text
                # print(catalog_0)
                catalog.append(catalog_0)

    # 提取目录以及页码
    c_page = []
    for jj in catalog:
        # print(jj)
        with fitz.open(output_pdf) as doc:
            for index, page in enumerate(doc, start=1):
                if jj in page.getText():
                    c_page_0 = jj + f'{index}'
                    # print(c_page_0)
                    c_page.append(c_page_0)

    # 目录页码列表
    page = []
    for kk in c_page:
        page_0 = kk.replace(kk.rstrip(string.digits), '')
        page.append(page_0)

    # 增加pdf标签
    output = PdfFileWriter()
    f = open(output_pdf, 'rb')
    input1 = PdfFileReader(f)
    pageCount = input1.getNumPages()
    for iPage in range(pageCount):
        output.addPage(input1.getPage(iPage))

    for l in range(len(catalog)):
        levels_num = catalog[l].count('.')
        output.addBookmark((2 * levels_num - 1) * ' ' + catalog[l], int(page[l]) - 1)  # 添加父书签
        # output.addBookmark('Hello, World',0,parent)    #添加子书签
    with open(pdf_marks_path, 'wb') as fout:
        output.write(fout)
    f.close()

def veri_report_word(data_path,doc_path):
    pass

def veri_report_word(data_path,pdf_path):
    pass

def word_to_pdf(doc_path,pdf_path):
    pass

if __name__ == '__main__':
    data_path = r'H:\task\project\department\202004-检验系统研发\公报转pdf测试数据'  # 根目录路径
    doc_path = r'H:\task\project\department\202004-检验系统研发\doc_test01.doc'  # 生成word文件路径
    pdf_path = r'H:\task\project\department\202004-检验系统研发\pdf_test01.pdf'  # 生成word转pdf文件,pdf不带标签
    pdf_marks_path = 'pdf_marks01.pdf'  # 生成带标签的pdf文件
    pdf_withmarks(data_path, doc_path, pdf_path, pdf_marks_path)


    
                
    
