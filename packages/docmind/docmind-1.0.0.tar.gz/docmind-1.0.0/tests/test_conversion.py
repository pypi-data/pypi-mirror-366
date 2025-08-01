#!/usr/bin/env python3
"""Test DocMind conversion with available dependencies."""

import sys
from pathlib import Path

def test_pdf_basic():
    """Test basic PDF processing without external tools."""
    try:
        import fitz
        
        pdf_path = Path("test-scenarios/input.pdf")
        if not pdf_path.exists():
            print("❌ PDF test file not found")
            return False
            
        print(f"📄 Testing PDF: {pdf_path}")
        
        # Test basic PDF opening
        doc = fitz.open(pdf_path)
        print(f"✅ PDF opened: {doc.page_count} pages")
        
        # Test basic text extraction
        if doc.page_count > 0:
            page = doc[0]
            text = page.get_text()
            print(f"✅ Text extracted: {len(text)} characters")
            
            # Test image extraction
            images = page.get_images()
            print(f"✅ Images found: {len(images)} images on first page")
        
        doc.close()
        return True
        
    except ImportError:
        print("❌ PyMuPDF not available")
        return False
    except Exception as e:
        print(f"❌ PDF test failed: {e}")
        return False

def test_docx_basic():
    """Test basic DOCX processing."""
    try:
        from docx import Document
        
        docx_path = Path("test-scenarios/input.docx")
        if not docx_path.exists():
            print("❌ DOCX test file not found")
            return False
            
        print(f"📄 Testing DOCX: {docx_path}")
        
        # Test basic DOCX opening
        doc = Document(docx_path)
        print(f"✅ DOCX opened: {len(doc.paragraphs)} paragraphs")
        
        # Test text extraction
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        
        print(f"✅ Text extracted: {len(text_content)} non-empty paragraphs")
        
        # Test table detection
        print(f"✅ Tables found: {len(doc.tables)} tables")
        
        return True
        
    except ImportError:
        print("❌ python-docx not available")
        return False
    except Exception as e:
        print(f"❌ DOCX test failed: {e}")
        return False

def test_image_processing():
    """Test image processing capabilities."""
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        import io
        
        print("🖼️  Testing image processing...")
        
        # Create a test image
        test_image = Image.new('RGB', (100, 100), 'white')
        
        # Test LLM optimizations
        enhancer = ImageEnhance.Contrast(test_image)
        enhanced = enhancer.enhance(1.1)
        
        enhancer = ImageEnhance.Sharpness(enhanced)
        sharpened = enhancer.enhance(1.2)
        
        # Test unsharp mask
        final = sharpened.filter(ImageFilter.UnsharpMask(radius=1, percent=120, threshold=3))
        
        print("✅ Image processing works")
        return True
        
    except ImportError:
        print("❌ PIL not available")
        return False
    except Exception as e:
        print(f"❌ Image processing failed: {e}")
        return False

def test_config_system():
    """Test configuration system."""
    try:
        from docmind.config import DocMindConfig, get_preset_config, QUALITY_PRESETS
        
        print("⚙️  Testing configuration system...")
        
        # Test default config
        config = DocMindConfig()
        print(f"✅ Default config created")
        
        # Test presets
        for preset_name in QUALITY_PRESETS:
            preset_config = get_preset_config(preset_name)
            print(f"✅ Preset '{preset_name}' loaded")
        
        return True
        
    except Exception as e:
        print(f"❌ Config system failed: {e}")
        return False

def main():
    """Run all tests."""
    print("🧪 Testing DocMind Components")
    print("=" * 40)
    
    tests = [
        ("PDF Processing", test_pdf_basic),
        ("DOCX Processing", test_docx_basic),
        ("Image Processing", test_image_processing),
        ("Configuration System", test_config_system),
    ]
    
    passed = 0
    total = len(tests)
    
    for test_name, test_func in tests:
        print(f"\n🔍 {test_name}:")
        try:
            if test_func():
                passed += 1
        except Exception as e:
            print(f"❌ Test {test_name} crashed: {e}")
    
    print(f"\n📊 Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All component tests passed!")
        print("\n💡 To test full conversion, install external tools:")
        print("   brew install pandoc tesseract imagemagick")
    else:
        print("⚠️  Some tests failed. Check dependencies.")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)