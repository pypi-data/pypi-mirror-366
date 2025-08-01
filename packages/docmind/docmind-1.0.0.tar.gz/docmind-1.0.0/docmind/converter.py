#!/usr/bin/env python3
"""DocMind: Universal AI-optimized document converter."""

import subprocess
import hashlib
import pickle
import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import logging
from datetime import datetime

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    from PIL import Image, ImageEnhance, ImageFilter

from .config import DocMindConfig
from .exceptions import *

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocMind:
    """
    DocMind: Universal AI-optimized document converter for technical documents.
    """
    
    def __init__(self, config: Optional[DocMindConfig] = None, output_dir: str = "output"):
        """Initialize DocMind with configuration."""
        self.config = config or DocMindConfig()
        self.output_dir = Path(output_dir)
        self.media_dir = self.output_dir / "media"
        self.tables_dir = self.output_dir / "tables"
        self.cache_dir = self.output_dir / ".cache"
        
        # Statistics
        self.stats = {
            'pages_processed': 0,
            'images_extracted': 0,
            'tables_extracted': 0,
            'ocr_pages': 0,
            'errors': 0,
            'start_time': None,
            'end_time': None
        }
        
        # Check tool availability
        self.tools = self._check_tools()
        
        # Setup directories
        self._setup_directories()
        
    def _setup_directories(self) -> None:
        """Create necessary directories."""
        for directory in [self.output_dir, self.media_dir, self.tables_dir]:
            directory.mkdir(parents=True, exist_ok=True)
            
        if self.config.performance.cache_enabled:
            self.cache_dir.mkdir(parents=True, exist_ok=True)
    
    def _check_tools(self) -> Dict[str, bool]:
        """Check availability of external tools with detailed information."""
        tools = {}
        
        # Essential tools
        for tool, cmd in [
            ('pandoc', ['pandoc', '--version']),
            ('imagemagick', ['magick', '-version']),
            ('tesseract', ['tesseract', '--version']),
            ('libreoffice', ['soffice', '--version'])
        ]:
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                tools[tool] = result.returncode == 0
                if tools[tool]:
                    logger.info(f"✅ {tool.capitalize()} found")
                else:
                    logger.warning(f"⚠️  {tool.capitalize()} found but returned error")
            except (FileNotFoundError, subprocess.TimeoutExpired):
                tools[tool] = False
                logger.warning(f"❌ {tool.capitalize()} not found")
        
        # Python libraries
        tools['tqdm'] = TQDM_AVAILABLE
        tools['tesseract_py'] = TESSERACT_AVAILABLE
        
        return tools
    
    def validate_input(self, file_path: Path) -> None:
        """Validate input file before processing."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.performance.max_file_size_mb:
            raise FileTooLargeError(
                f"File too large: {file_size_mb:.1f}MB. "
                f"Maximum allowed: {self.config.performance.max_file_size_mb}MB"
            )
        
        # Check file format
        suffix = file_path.suffix.lower()
        if suffix not in ['.pdf', '.docx']:
            raise UnsupportedFormatError(f"Unsupported format: {suffix}")
        
        # Basic corruption check
        try:
            if suffix == '.pdf':
                import fitz
                doc = fitz.open(file_path)
                if doc.page_count == 0:
                    raise CorruptedFileError("PDF has no pages")
                doc.close()
            elif suffix == '.docx':
                from docx import Document
                doc = Document(file_path)
                if not doc.paragraphs and not doc.tables:
                    raise CorruptedFileError("DOCX appears to be empty")
        except Exception as e:
            if isinstance(e, (CorruptedFileError, UnsupportedFormatError, FileTooLargeError)):
                raise
            raise CorruptedFileError(f"File appears corrupted: {e}")
    
    def get_file_hash(self, file_path: Path) -> str:
        """Get SHA256 hash of file for caching."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()[:16]  # Use first 16 chars
    
    def load_progress(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Load conversion progress from cache."""
        if not self.config.performance.enable_resume:
            return None
        
        file_hash = self.get_file_hash(file_path)
        progress_file = self.cache_dir / f"{file_hash}_progress.pkl"
        
        if progress_file.exists():
            try:
                with open(progress_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                logger.warning(f"Could not load progress: {e}")
        
        return None
    
    def save_progress(self, file_path: Path, progress: Dict[str, Any]) -> None:
        """Save conversion progress to cache."""
        if not self.config.performance.enable_resume:
            return
        
        file_hash = self.get_file_hash(file_path)
        progress_file = self.cache_dir / f"{file_hash}_progress.pkl"
        
        try:
            with open(progress_file, 'wb') as f:
                pickle.dump(progress, f)
        except Exception as e:
            logger.warning(f"Could not save progress: {e}")
    
    def convert_pdf(self, pdf_path: Path) -> bool:
        """Convert PDF with all features."""
        try:
            import fitz
        except ImportError:
            raise DependencyMissingError("PyMuPDF (fitz) required for PDF processing")
        
        logger.info(f"🔄 Processing PDF: {pdf_path}")
        
        # Load progress if resuming
        progress = self.load_progress(pdf_path)
        start_page = 0
        if progress:
            start_page = progress.get('last_page', 0)
            logger.info(f"📄 Resuming from page {start_page + 1}")
        
        doc = fitz.open(pdf_path)
        total_pages = doc.page_count
        
        # Initialize progress bar
        pbar = None
        if TQDM_AVAILABLE:
            pbar = tqdm(total=total_pages, desc="Converting PDF", initial=start_page)
        
        md_output = []
        images_extracted = 0
        tables_extracted = 0
        
        try:
            # Process pages in chunks for memory efficiency
            chunk_size = self.config.performance.chunk_size_pages
            
            for chunk_start in range(start_page, total_pages, chunk_size):
                chunk_end = min(chunk_start + chunk_size, total_pages)
                
                # Process chunk
                chunk_result = self._process_pdf_chunk(
                    doc, chunk_start, chunk_end, pdf_path.stem
                )
                
                md_output.extend(chunk_result['markdown'])
                images_extracted += chunk_result['images']
                tables_extracted += chunk_result['tables']
                
                # Update progress
                if pbar:
                    pbar.update(chunk_end - chunk_start)
                
                # Save progress
                self.save_progress(pdf_path, {
                    'last_page': chunk_end - 1,
                    'images_extracted': images_extracted,
                    'tables_extracted': tables_extracted
                })
                
                self.stats['pages_processed'] = chunk_end
        
        finally:
            doc.close()
            if pbar:
                pbar.close()
        
        # Save final markdown
        self._save_markdown(md_output, images_extracted, tables_extracted)
        
        # Update stats
        self.stats['images_extracted'] = images_extracted
        self.stats['tables_extracted'] = tables_extracted
        
        logger.info(f"✅ PDF processed: {total_pages} pages, {images_extracted} images, {tables_extracted} tables")
        return True
    
    def _process_pdf_chunk(self, doc, start_page: int, end_page: int, doc_name: str) -> Dict[str, Any]:
        """Process a chunk of PDF pages."""
        md_output = []
        images_extracted = 0
        tables_extracted = 0
        
        for page_num in range(start_page, end_page):
            try:
                page = doc[page_num]
                
                # Add page header
                md_output.append(f"# Page {page_num + 1}\n")
                
                # Extract text with Pandoc if available, fallback to direct extraction
                if self.tools['pandoc']:
                    text = self._extract_text_pandoc(page)
                else:
                    text = page.get_text("markdown")
                
                md_output.append(text)
                
                # Extract images
                images_on_page = self._extract_page_images(page, page_num + 1, doc_name)
                images_extracted += len(images_on_page)
                
                # Add image references
                for img_filename in images_on_page:
                    md_output.append(f"![Image](media/{img_filename})\n")
                
                # Extract tables if enabled
                if self.config.conversion.extract_tables:
                    tables_on_page = self._extract_page_tables(page, page_num + 1)
                    tables_extracted += len(tables_on_page)
                
                # OCR if needed
                if self.config.conversion.ocr_enabled and self._needs_ocr(page):
                    ocr_text = self._perform_ocr(page, page_num + 1)
                    if ocr_text:
                        md_output.append(f"\n## OCR Content (Page {page_num + 1})\n")
                        md_output.append(ocr_text)
                        self.stats['ocr_pages'] += 1
                
                md_output.append("\n---\n")
                
            except Exception as e:
                logger.error(f"Error processing page {page_num + 1}: {e}")
                self.stats['errors'] += 1
                md_output.append(f"\n*[Error processing page {page_num + 1}: {e}]*\n")
        
        return {
            'markdown': md_output,
            'images': images_extracted,
            'tables': tables_extracted
        }
    
    def _extract_text_pandoc(self, page) -> str:
        """Extract text using Pandoc for better formatting."""
        try:
            # Get HTML text from page
            html_text = page.get_text("html")
            
            # Use Pandoc to convert HTML to Markdown
            cmd = ['pandoc', '-f', 'html', '-t', 'markdown'] + self.config.conversion.pandoc_extra_args
            
            result = subprocess.run(
                cmd, input=html_text, capture_output=True, text=True, encoding='utf-8'
            )
            
            if result.returncode == 0:
                return result.stdout
            else:
                logger.warning(f"Pandoc conversion failed: {result.stderr}")
                return page.get_text("markdown")
                
        except Exception as e:
            logger.warning(f"Pandoc text extraction failed: {e}")
            return page.get_text("markdown")
    
    def _extract_page_images(self, page, page_num: int, doc_name: str) -> List[str]:
        """Extract and optimize images from a page."""
        image_files = []
        
        for img_index, img in enumerate(page.get_images(full=True)):
            try:
                # Extract image
                xref = img[0]
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image.get("ext", "").lower()
                
                # Generate filename
                img_filename = f"{doc_name}_p{page_num:03d}_img{img_index:02d}.png"
                img_path = self.media_dir / img_filename
                
                # Process and optimize image
                if self._process_and_save_image(image_bytes, img_path, image_ext):
                    image_files.append(img_filename)
                
            except Exception as e:
                logger.warning(f"Failed to extract image {img_index} from page {page_num}: {e}")
        
        return image_files
    
    def _process_and_save_image(self, image_bytes: bytes, output_path: Path, original_ext: str) -> bool:
        """Process and save image with LLM optimization."""
        try:
            # Open image
            image = Image.open(io.BytesIO(image_bytes))
            
            # Convert transparency to white background
            if image.mode in ['RGBA', 'P', 'LA']:
                rgb_image = Image.new('RGB', image.size, 'white')
                if image.mode in ['RGBA', 'LA']:
                    rgb_image.paste(image, mask=image.split()[-1])
                else:
                    rgb_image.paste(image)
                image = rgb_image
            
            # Apply LLM optimizations if enabled
            if self.config.image.optimize_for_llm:
                image = self._optimize_image_for_llm(image)
            
            # Ensure minimum resolution
            if (image.width < self.config.image.min_width or 
                image.height < self.config.image.min_height):
                
                scale_w = self.config.image.min_width / image.width
                scale_h = self.config.image.min_height / image.height
                scale = max(scale_w, scale_h, 1.0)  # Don't downscale
                
                new_size = (int(image.width * scale), int(image.height * scale))
                image = image.resize(new_size, Image.LANCZOS)
            
            # Save with quality settings
            save_kwargs = {
                'format': self.config.image.format,
                'quality': self.config.image.quality,
                'optimize': True
            }
            
            image.save(output_path, **save_kwargs)
            return True
            
        except Exception as e:
            logger.warning(f"Failed to process image: {e}")
            return False
    
    def _optimize_image_for_llm(self, image: Image.Image) -> Image.Image:
        """Apply LLM-specific optimizations to image."""
        # Enhance contrast
        if self.config.image.enhance_contrast != 1.0:
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(self.config.image.enhance_contrast)
        
        # Enhance sharpness  
        if self.config.image.enhance_sharpness != 1.0:
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(self.config.image.enhance_sharpness)
        
        # Apply unsharp mask for better text/diagram clarity
        image = image.filter(ImageFilter.UnsharpMask(radius=1, percent=120, threshold=3))
        
        return image
    
    def _extract_page_tables(self, page, page_num: int) -> List[str]:
        """Extract tables from page using pdfplumber."""
        tables = []
        
        try:
            # This would require integrating with pdfplumber
            # For now, placeholder implementation
            logger.info(f"Table extraction for page {page_num} - feature pending")
            
        except Exception as e:
            logger.warning(f"Table extraction failed for page {page_num}: {e}")
        
        return tables
    
    def _needs_ocr(self, page) -> bool:
        """Determine if page needs OCR (has images but little text)."""
        if not self.config.conversion.ocr_enabled or not TESSERACT_AVAILABLE:
            return False
        
        text_length = len(page.get_text().strip())
        image_count = len(page.get_images())
        
        # Heuristic: OCR if lots of images but little text
        return image_count > 0 and text_length < 100
    
    def _perform_ocr(self, page, page_num: int) -> Optional[str]:
        """Perform OCR on page."""
        if not TESSERACT_AVAILABLE:
            return None
        
        try:
            import fitz
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for better OCR
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(
                image, 
                lang=self.config.conversion.ocr_language,
                config='--oem 3 --psm 6'
            )
            
            return ocr_text.strip()
            
        except Exception as e:
            logger.warning(f"OCR failed for page {page_num}: {e}")
            return None
    
    def convert_docx(self, docx_path: Path) -> bool:
        """Convert DOCX with all features."""
        try:
            from docx import Document
        except ImportError:
            raise DependencyMissingError("python-docx required for DOCX processing")
        
        logger.info(f"🔄 Processing DOCX: {docx_path}")
        
        doc = Document(docx_path)
        md_output = []
        
        # Process paragraphs and tables
        total_elements = len(doc.paragraphs) + len(doc.tables)
        
        pbar = None
        if TQDM_AVAILABLE:
            pbar = tqdm(total=total_elements, desc="Converting DOCX")
        
        try:
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    if para.style.name.startswith('Heading'):
                        level = self._extract_heading_level(para.style.name)
                        md_output.append(f"{'#' * level} {para.text.strip()}\n")
                    else:
                        md_output.append(f"{para.text.strip()}\n\n")
                
                if pbar:
                    pbar.update(1)
            
            # Process tables
            tables_extracted = 0
            for table_idx, table in enumerate(doc.tables):
                table_md = self._convert_docx_table(table, table_idx)
                if table_md:
                    md_output.append(table_md)
                    tables_extracted += 1
                
                if pbar:
                    pbar.update(1)
            
            # Extract embedded images
            images_extracted = self._extract_docx_images(docx_path)
            
        finally:
            if pbar:
                pbar.close()
        
        # Save results
        self._save_markdown(md_output, images_extracted, tables_extracted)
        
        # Update stats
        self.stats['images_extracted'] = images_extracted
        self.stats['tables_extracted'] = tables_extracted
        
        logger.info(f"✅ DOCX processed: {images_extracted} images, {tables_extracted} tables")
        return True
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from DOCX style name."""
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1
    
    def _convert_docx_table(self, table, table_idx: int) -> str:
        """Convert DOCX table to Markdown."""
        try:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if not table_data:
                return ""
            
            # Create markdown table
            md_table = []
            
            # Header row
            header = table_data[0]
            md_table.append("| " + " | ".join(header) + " |")
            md_table.append("| " + " | ".join(["---"] * len(header)) + " |")
            
            # Data rows
            for row in table_data[1:]:
                padded_row = row + [""] * (len(header) - len(row))
                md_table.append("| " + " | ".join(padded_row[:len(header)]) + " |")
            
            # Save separate table file if configured
            if self.config.output.separate_tables:
                table_file = self.tables_dir / f"docx_table_{table_idx:02d}.md"
                with open(table_file, 'w', encoding='utf-8') as f:
                    f.write("\n".join(md_table))
            
            return "\n" + "\n".join(md_table) + "\n\n"
            
        except Exception as e:
            logger.warning(f"Failed to convert table {table_idx}: {e}")
            return ""
    
    def _extract_docx_images(self, docx_path: Path) -> int:
        """Extract images from DOCX file."""
        try:
            import zipfile
            
            images_extracted = 0
            
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                image_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for idx, img_file in enumerate(image_files):
                    try:
                        img_data = zip_file.read(img_file)
                        img_name = f"docx_img_{idx:03d}.png"
                        img_path = self.media_dir / img_name
                        
                        if self._process_and_save_image(img_data, img_path, 
                                                      Path(img_file).suffix.lower()):
                            images_extracted += 1
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract image {img_file}: {e}")
            
            return images_extracted
            
        except Exception as e:
            logger.warning(f"DOCX image extraction failed: {e}")
            return 0
    
    def _save_markdown(self, md_content: List[str], images_count: int, tables_count: int) -> None:
        """Save final markdown with metadata and enhancements."""
        output_file = self.output_dir / "output.md"
        
        # Add metadata header if configured
        content = []
        
        if self.config.conversion.include_metadata:
            content.extend([
                "---",
                f"title: \"Document Conversion\"",
                f"generated_by: \"DocMind v1.0.0\"",
                f"generated_at: \"{datetime.now().isoformat()}\"",
                f"images_extracted: {images_count}",
                f"tables_extracted: {tables_count}",
                f"pages_processed: {self.stats['pages_processed']}",
                "---",
                ""
            ])
        
        # Add table of contents if configured
        if self.config.output.include_toc:
            content.extend([
                "# Table of Contents",
                "",
                "- [Document Content](#document-content)",
                f"- [Images ({images_count})](#images)" if images_count > 0 else "",
                f"- [Tables ({tables_count})](#tables)" if tables_count > 0 else "",
                "",
                "# Document Content",
                ""
            ])
        
        # Add main content
        content.extend(md_content)
        
        # Add appendices
        if images_count > 0 and self.config.output.separate_images:
            content.extend([
                "",
                "# Images",
                "",
                f"This document contains {images_count} extracted images in the `media/` directory.",
                ""
            ])
        
        if tables_count > 0 and self.config.output.separate_tables:
            content.extend([
                "",
                "# Tables", 
                "",
                f"This document contains {tables_count} extracted tables in the `tables/` directory.",
                ""
            ])
        
        # Write final content
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("\n".join(content))
        
        logger.info(f"📄 Markdown saved: {output_file}")
    
    def convert(self, input_files: Union[Path, List[Path]]) -> bool:
        """Main conversion method supporting single files or batch processing."""
        self.stats['start_time'] = datetime.now()
        
        # Normalize input to list
        if isinstance(input_files, (str, Path)):
            input_files = [Path(input_files)]
        else:
            input_files = [Path(f) for f in input_files]
        
        logger.info(f"🚀 Starting conversion of {len(input_files)} file(s)")
        
        success_count = 0
        
        for file_path in input_files:
            try:
                # Validate input
                self.validate_input(file_path)
                
                # Determine conversion method
                suffix = file_path.suffix.lower()
                
                if suffix == '.pdf':
                    success = self.convert_pdf(file_path)
                elif suffix == '.docx':
                    success = self.convert_docx(file_path)
                else:
                    raise UnsupportedFormatError(f"Unsupported format: {suffix}")
                
                if success:
                    success_count += 1
                
            except Exception as e:
                logger.error(f"❌ Failed to convert {file_path}: {e}")
                self.stats['errors'] += 1
        
        self.stats['end_time'] = datetime.now()
        
        # Print final statistics
        self._print_statistics(success_count, len(input_files))
        
        return success_count == len(input_files)
    
    def _print_statistics(self, success_count: int, total_count: int) -> None:
        """Print conversion statistics."""
        duration = self.stats['end_time'] - self.stats['start_time']
        
        logger.info("\n" + "="*50)
        logger.info("📊 CONVERSION STATISTICS")
        logger.info("="*50)
        logger.info(f"Files processed: {success_count}/{total_count}")
        logger.info(f"Pages processed: {self.stats['pages_processed']}")
        logger.info(f"Images extracted: {self.stats['images_extracted']}")
        logger.info(f"Tables extracted: {self.stats['tables_extracted']}")
        logger.info(f"OCR pages: {self.stats['ocr_pages']}")
        logger.info(f"Errors encountered: {self.stats['errors']}")
        logger.info(f"Total time: {duration}")
        logger.info(f"Output directory: {self.output_dir}")
        logger.info("="*50)

def main():
    """Simple CLI entry point for basic usage."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="DocMind: AI-optimized converter for technical documents"
    )
    parser.add_argument("input_file", help="Path to input PDF or DOCX file")
    parser.add_argument("-o", "--output", default="output", 
                       help="Output directory (default: output)")
    
    args = parser.parse_args()
    
    try:
        converter = DocMind(output_dir=args.output)
        success = converter.convert(args.input_file)
        
        if success:
            print("✅ Conversion completed successfully!")
        else:
            print("❌ Conversion failed")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())