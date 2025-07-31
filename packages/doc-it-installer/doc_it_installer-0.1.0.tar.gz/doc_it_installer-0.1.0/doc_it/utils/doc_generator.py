# doc_it/utils/doc_generator.py

from rich.console import Console
from docx import Document
from pylatex import Document as LaTexDocument, Section, Subsection, Command
from pylatex.utils import NoEscape

# Import necessary functions from our other modules
from doc_it.utils.config_handler import PROJECT_ROOT, get_config
from doc_it.db.database import get_codebase_summary, get_explanations

console = Console()

def generate_markdown() -> str | None:
    """
    Fetches all data from the database and compiles it into a single
    Markdown documentation file.
    """
    config = get_config()
    summary = get_codebase_summary()
    explanations = get_explanations()

    if not summary and not explanations:
        return None

    md_content = f"# Doc-It: AI-Generated Code Documentation\n\n"
    md_content += f"**Author:** {config.get('author', 'N/A')}\n\n---\n\n"
    if summary:
        md_content += f"## 1. Codebase Overview\n\n{summary}\n\n---\n\n"
    if explanations:
        md_content += "## 2. Change History & Explanations\n\n"
        for entry in explanations:
            md_content += f"### File: `{entry['filepath']}`\n"
            md_content += f"*Timestamp: {entry['timestamp']}*\n\n"
            md_content += "#### AI-Generated Explanation:\n"
            md_content += f"> {entry['explanation']}\n\n"
    
    output_path = PROJECT_ROOT / "documentation.md"
    try:
        with open(output_path, "w", encoding="utf-8") as f: f.write(md_content)
        return str(output_path)
    except Exception as e:
        console.print(f"[bold red]Error writing Markdown file: {e}[/bold red]")
        return None

def generate_docx() -> str | None:
    """Generates a .docx documentation file."""
    config = get_config()
    summary = get_codebase_summary()
    explanations = get_explanations()

    if not summary and not explanations: return None

    doc = Document()
    doc.add_heading('Doc-It: AI-Generated Code Documentation', level=0)
    doc.add_paragraph(f"Author: {config.get('author', 'N/A')}")
    
    if summary:
        doc.add_heading('Codebase Overview', level=1)
        doc.add_paragraph(summary)

    if explanations:
        doc.add_heading('Change History & Explanations', level=1)
        for entry in explanations:
            doc.add_heading(f"File: {entry['filepath']}", level=2)
            doc.add_paragraph(f"Timestamp: {entry['timestamp']}").italic = True
            doc.add_paragraph("AI-Generated Explanation:", style='Intense Quote')
            doc.add_paragraph(entry['explanation'])

    output_path = PROJECT_ROOT / "documentation.docx"
    try:
        doc.save(output_path)
        return str(output_path)
    except Exception as e:
        console.print(f"[bold red]Error writing DOCX file: {e}[/bold red]")
        return None

def generate_latex() -> str | None:
    """Generates a LaTeX documentation file."""
    config = get_config()
    summary = get_codebase_summary()
    explanations = get_explanations()

    if not summary and not explanations: return None

    doc = LaTexDocument()
    doc.preamble.append(Command('title', 'Doc-It: AI-Generated Code Documentation'))
    doc.preamble.append(Command('author', config.get('author', 'N/A')))
    doc.preamble.append(Command('date', NoEscape(r'\today')))
    doc.append(NoEscape(r'\maketitle'))

    if summary:
        with doc.create(Section('Codebase Overview')):
            doc.append(summary)

    if explanations:
        with doc.create(Section('Change History & Explanations')):
            for entry in explanations:
                with doc.create(Subsection(f"File: {entry['filepath']}")):
                    doc.append(f"Timestamp: {entry['timestamp']}\n\n")
                    doc.append(NoEscape(r'\textbf{AI-Generated Explanation:}\par'))
                    doc.append(entry['explanation'])

    output_path = PROJECT_ROOT / "documentation" # .tex is added automatically
    try:
        doc.generate_tex(str(output_path))
        return f"{output_path}.tex"
    except Exception as e:
        console.print(f"[bold red]Error writing LaTeX file: {e}[/bold red]")
        return None
