"""
AI Helper Agent - Enhanced File Input Handler
Supports reading various file types including .py, .txt, .md, .json, .csv, etc.
"""

import os
import pathlib
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
import mimetypes
import json
import csv
import sqlite3

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

import structlog
logger = structlog.get_logger()


class EnhancedFileHandler:
    """Enhanced file handler for various file types with AI integration"""
    
    SUPPORTED_EXTENSIONS = {
        # Text files
        '.txt': 'text',
        '.md': 'markdown', 
        '.rst': 'text',
        '.log': 'text',
        
        # Code files
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.html': 'html',
        '.css': 'css',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.cs': 'csharp',
        '.php': 'php',
        '.rb': 'ruby',
        '.go': 'go',
        '.rs': 'rust',
        '.sql': 'sql',
        '.sh': 'bash',
        '.bat': 'batch',
        '.ps1': 'powershell',
        
        # Config files
        '.json': 'json',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.toml': 'toml',
        '.ini': 'ini',
        '.cfg': 'config',
        '.env': 'env',
        '.conf': 'config',
        
        # Data files
        '.csv': 'csv',
        '.tsv': 'tsv',
        '.xml': 'xml',
        
        # Documents
        '.docx': 'docx',
        '.pdf': 'pdf',
        
        # Other
        '.sqlite': 'sqlite',
        '.db': 'sqlite'
    }
    
    def __init__(self):
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.encoding_attempts = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
        
    def is_supported_file(self, file_path: str) -> bool:
        """Check if file type is supported"""
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
        
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        path = Path(file_path)
        
        if not path.exists():
            return {"error": "File not found"}
            
        try:
            stat = path.stat()
            mime_type, _ = mimetypes.guess_type(str(path))
            
            info = {
                "name": path.name,
                "path": str(path.absolute()),
                "size": stat.st_size,
                "size_human": self._format_file_size(stat.st_size),
                "extension": path.suffix.lower(),
                "mime_type": mime_type,
                "file_type": self.SUPPORTED_EXTENSIONS.get(path.suffix.lower(), "unknown"),
                "is_supported": self.is_supported_file(file_path),
                "modified_time": stat.st_mtime,
                "is_binary": self._is_binary_file(path)
            }
            
            return info
            
        except Exception as e:
            logger.error("Error getting file info", file_path=file_path, error=str(e))
            return {"error": str(e)}
    
    def read_file_content(self, file_path: str) -> Dict[str, Any]:
        """Read file content with appropriate handler based on file type"""
        file_info = self.get_file_info(file_path)
        
        if "error" in file_info:
            return file_info
            
        if file_info["size"] > self.max_file_size:
            return {"error": f"File too large. Maximum size: {self._format_file_size(self.max_file_size)}"}
            
        file_type = file_info["file_type"]
        
        try:
            # Route to appropriate handler
            if file_type in ['text', 'markdown', 'python', 'javascript', 'typescript', 'html', 'css', 
                           'java', 'cpp', 'c', 'csharp', 'php', 'ruby', 'go', 'rust', 'sql', 
                           'bash', 'batch', 'powershell', 'config', 'env']:
                return self._read_text_file(file_path, file_info)
                
            elif file_type == 'json':
                return self._read_json_file(file_path, file_info)
                
            elif file_type in ['yaml', 'yml']:
                return self._read_yaml_file(file_path, file_info)
                
            elif file_type == 'csv':
                return self._read_csv_file(file_path, file_info)
                
            elif file_type == 'docx':
                return self._read_docx_file(file_path, file_info)
                
            elif file_type == 'pdf':
                return self._read_pdf_file(file_path, file_info)
                
            elif file_type == 'sqlite':
                return self._read_sqlite_file(file_path, file_info)
                
            else:
                return self._read_text_file(file_path, file_info)  # Fallback to text
                
        except Exception as e:
            logger.error("Error reading file", file_path=file_path, error=str(e))
            return {"error": f"Failed to read file: {str(e)}", "file_info": file_info}
    
    def _read_text_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read text-based files with encoding detection"""
        path = Path(file_path)
        
        # Try different encodings
        for encoding in self.encoding_attempts:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                    
                return {
                    "content": content,
                    "file_info": file_info,
                    "encoding": encoding,
                    "lines": len(content.split('\n')),
                    "characters": len(content),
                    "words": len(content.split()) if content else 0
                }
                
            except UnicodeDecodeError:
                continue
                
        return {"error": "Could not decode file with any supported encoding"}
    
    def _read_json_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read JSON files with validation"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            return {
                "content": json.dumps(data, indent=2),
                "parsed_data": data,
                "file_info": file_info,
                "json_valid": True,
                "data_type": type(data).__name__,
                "keys": list(data.keys()) if isinstance(data, dict) else None,
                "length": len(data) if isinstance(data, (list, dict)) else None
            }
            
        except json.JSONDecodeError as e:
            # Fallback to reading as text
            text_result = self._read_text_file(file_path, file_info)
            text_result["json_error"] = str(e)
            text_result["json_valid"] = False
            return text_result
    
    def _read_yaml_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read YAML files"""
        if not YAML_AVAILABLE:
            return self._read_text_file(file_path, file_info)
            
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
                
            return {
                "content": yaml.dump(data, default_flow_style=False, indent=2),
                "parsed_data": data,
                "file_info": file_info,  
                "yaml_valid": True,
                "data_type": type(data).__name__
            }
            
        except yaml.YAMLError as e:
            text_result = self._read_text_file(file_path, file_info)
            text_result["yaml_error"] = str(e)
            text_result["yaml_valid"] = False
            return text_result
    
    def _read_csv_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read CSV files with analysis"""
        try:
            # Try to detect delimiter
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
            # Read CSV
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)
                
            # Generate preview
            preview_rows = rows[:10]  # First 10 rows
            
            return {
                "content": self._format_csv_preview(rows[:100]),  # First 100 rows as text
                "parsed_data": rows,
                "file_info": file_info,
                "csv_info": {
                    "columns": reader.fieldnames,
                    "row_count": len(rows),
                    "column_count": len(reader.fieldnames) if reader.fieldnames else 0,
                    "delimiter": delimiter,
                    "preview": preview_rows
                }
            }
            
        except Exception as e:
            return self._read_text_file(file_path, file_info)
    
    def _read_docx_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read DOCX files"""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx package not installed"}
            
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n'.join(paragraphs)
            
            return {
                "content": content,
                "file_info": file_info,
                "document_info": {
                    "paragraphs": len(doc.paragraphs),
                    "non_empty_paragraphs": len(paragraphs),
                    "characters": len(content),
                    "words": len(content.split()) if content else 0
                }
            }
            
        except Exception as e:
            return {"error": f"Failed to read DOCX: {str(e)}"}
    
    def _read_pdf_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read PDF files"""
        if not PDF_AVAILABLE:
            return {"error": "PyPDF2 package not installed"}
            
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        pages.append(f"--- Page {i+1} ---\n{text}")
                
                content = '\n\n'.join(pages)
                
            return {
                "content": content,
                "file_info": file_info,
                "pdf_info": {
                    "pages": len(reader.pages),
                    "pages_with_text": len(pages),
                    "characters": len(content)
                }
            }
            
        except Exception as e:
            return {"error": f"Failed to read PDF: {str(e)}"}
    
    def _read_sqlite_file(self, file_path: str, file_info: Dict) -> Dict[str, Any]:
        """Read SQLite database structure"""
        try:
            conn = sqlite3.connect(file_path)
            cursor = conn.cursor()
            
            # Get table names
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = [row[0] for row in cursor.fetchall()]
            
            db_info = {"tables": {}}
            
            for table in tables:
                # Get table info
                cursor.execute(f"PRAGMA table_info({table})")
                columns = cursor.fetchall()
                
                cursor.execute(f"SELECT COUNT(*) FROM {table}")
                row_count = cursor.fetchone()[0]
                
                db_info["tables"][table] = {
                    "columns": [{"name": col[1], "type": col[2]} for col in columns],
                    "row_count": row_count
                }
            
            conn.close()
            
            # Generate content summary
            content = self._format_sqlite_summary(db_info)
            
            return {
                "content": content,
                "parsed_data": db_info,
                "file_info": file_info,
                "database_info": {
                    "tables": len(tables),
                    "total_rows": sum(t["row_count"] for t in db_info["tables"].values())
                }
            }
            
        except Exception as e:
            return {"error": f"Failed to read SQLite database: {str(e)}"}
    
    def _format_csv_preview(self, rows: List[Dict]) -> str:
        """Format CSV data for preview"""
        if not rows:
            return "Empty CSV file"
            
        # Create a simple table format
        lines = []
        if rows:
            # Header
            headers = list(rows[0].keys())
            lines.append(" | ".join(headers))
            lines.append("-" * len(lines[0]))
            
            # Data rows
            for row in rows[:20]:  # First 20 rows
                values = [str(row.get(h, "")) for h in headers]
                lines.append(" | ".join(values))
                
        return "\n".join(lines)
    
    def _format_sqlite_summary(self, db_info: Dict) -> str:
        """Format SQLite database summary"""
        lines = ["SQLite Database Structure:\n"]
        
        for table_name, table_info in db_info["tables"].items():
            lines.append(f"Table: {table_name} ({table_info['row_count']} rows)")
            lines.append("Columns:")
            for col in table_info["columns"]:
                lines.append(f"  - {col['name']} ({col['type']})")
            lines.append("")
            
        return "\n".join(lines)
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"
    
    def _is_binary_file(self, path: Path) -> bool:
        """Check if file is binary"""
        try:
            with open(path, 'rb') as f:
                chunk = f.read(8192)
                return b'\0' in chunk
        except:
            return True
    
    def get_file_suggestions(self, directory: str = ".") -> List[Dict[str, Any]]:
        """Get suggestions for files in directory"""
        try:
            path = Path(directory)
            suggestions = []
            
            for file_path in path.rglob("*"):
                if file_path.is_file() and self.is_supported_file(str(file_path)):
                    info = self.get_file_info(str(file_path))
                    if "error" not in info:
                        suggestions.append({
                            "path": str(file_path),
                            "name": file_path.name,
                            "type": info["file_type"],
                            "size": info["size_human"],
                            "relative_path": str(file_path.relative_to(path))
                        })
                        
            return sorted(suggestions, key=lambda x: x["name"])
            
        except Exception as e:
            logger.error("Error getting file suggestions", error=str(e))
            return []


# Global file handler instance
file_handler = EnhancedFileHandler()
