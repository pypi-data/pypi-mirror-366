"""
Enhanced File Operations with OCR - Requirement #12
Provides comprehensive file reading capabilities including OCR for unreadable files
"""

import os
import re
import mimetypes
import logging
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any, List, Union, Tuple
import sqlite3
from datetime import datetime

# OCR and file processing dependencies
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    pdfplumber = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None

from .utils import get_user_data_dir


class FileOperationsDatabase:
    """Database manager for file operations history and OCR cache"""
    
    def __init__(self):
        self.user_data_dir = get_user_data_dir()
        self.db_path = self.user_data_dir / "file_operations.db"
        self._init_database()
    
    def _init_database(self):
        """Initialize file operations database tables"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # File processing history
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS file_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER,
                    processing_method TEXT NOT NULL,
                    ocr_used BOOLEAN DEFAULT FALSE,
                    success BOOLEAN NOT NULL,
                    error_message TEXT,
                    processing_time_ms INTEGER,
                    content_hash TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # OCR cache for processed images and PDFs
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ocr_cache (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT NOT NULL,
                    file_hash TEXT NOT NULL,
                    extracted_text TEXT,
                    confidence_score REAL,
                    ocr_engine TEXT DEFAULT 'tesseract',
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(file_path, file_hash)
                )
            """)
            
            # Supported file types registry
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS supported_formats (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    extension TEXT UNIQUE NOT NULL,
                    format_type TEXT NOT NULL,
                    processing_method TEXT NOT NULL,
                    requires_ocr BOOLEAN DEFAULT FALSE,
                    description TEXT
                )
            """)
            
            conn.commit()
            self._populate_supported_formats()
    
    def _populate_supported_formats(self):
        """Populate supported file formats"""
        formats = [
            # Code files
            ('.py', 'code', 'text', False, 'Python source code'),
            ('.c', 'code', 'text', False, 'C source code'),
            ('.cpp', 'code', 'text', False, 'C++ source code'),
            ('.java', 'code', 'text', False, 'Java source code'),
            ('.sol', 'code', 'text', False, 'Solidity smart contract'),
            ('.r', 'code', 'text', False, 'R statistical code'),
            ('.js', 'code', 'text', False, 'JavaScript code'),
            ('.ts', 'code', 'text', False, 'TypeScript code'),
            ('.html', 'code', 'text', False, 'HTML markup'),
            ('.css', 'code', 'text', False, 'CSS stylesheet'),
            ('.php', 'code', 'text', False, 'PHP code'),
            ('.rb', 'code', 'text', False, 'Ruby code'),
            ('.go', 'code', 'text', False, 'Go source code'),
            ('.rs', 'code', 'text', False, 'Rust source code'),
            
            # Text files
            ('.txt', 'text', 'text', False, 'Plain text file'),
            ('.md', 'text', 'text', False, 'Markdown document'),
            ('.rst', 'text', 'text', False, 'reStructuredText document'),
            ('.log', 'text', 'text', False, 'Log file'),
            ('.csv', 'data', 'structured', False, 'Comma-separated values'),
            ('.json', 'data', 'structured', False, 'JSON data'),
            ('.xml', 'data', 'structured', False, 'XML document'),
            ('.yaml', 'data', 'structured', False, 'YAML configuration'),
            ('.yml', 'data', 'structured', False, 'YAML configuration'),
            
            # Documents
            ('.pdf', 'document', 'pdf', True, 'PDF document'),
            ('.docx', 'document', 'docx', False, 'Microsoft Word document'),
            ('.doc', 'document', 'ocr', True, 'Legacy Word document'),
            ('.rtf', 'document', 'text', False, 'Rich Text Format'),
            
            # Images (for OCR)
            ('.png', 'image', 'ocr', True, 'PNG image'),
            ('.jpg', 'image', 'ocr', True, 'JPEG image'),
            ('.jpeg', 'image', 'ocr', True, 'JPEG image'),
            ('.gif', 'image', 'ocr', True, 'GIF image'),
            ('.bmp', 'image', 'ocr', True, 'Bitmap image'),
            ('.tiff', 'image', 'ocr', True, 'TIFF image'),
            ('.tif', 'image', 'ocr', True, 'TIFF image'),
        ]
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.executemany("""
                INSERT OR IGNORE INTO supported_formats 
                (extension, format_type, processing_method, requires_ocr, description)
                VALUES (?, ?, ?, ?, ?)
            """, formats)
            conn.commit()
    
    def log_file_operation(self, file_path: str, file_type: str, processing_method: str,
                          success: bool, ocr_used: bool = False, error_message: str = None,
                          processing_time_ms: int = None, file_size: int = None,
                          content_hash: str = None):
        """Log file processing operation"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO file_history 
                (file_path, file_type, file_size, processing_method, ocr_used, 
                 success, error_message, processing_time_ms, content_hash)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (file_path, file_type, file_size, processing_method, ocr_used,
                  success, error_message, processing_time_ms, content_hash))
            conn.commit()
    
    def cache_ocr_result(self, file_path: str, file_hash: str, extracted_text: str,
                        confidence_score: float = None, ocr_engine: str = 'tesseract'):
        """Cache OCR results for future use"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT OR REPLACE INTO ocr_cache 
                (file_path, file_hash, extracted_text, confidence_score, ocr_engine)
                VALUES (?, ?, ?, ?, ?)
            """, (file_path, file_hash, extracted_text, confidence_score, ocr_engine))
            conn.commit()
    
    def get_cached_ocr(self, file_path: str, file_hash: str) -> Optional[Dict[str, Any]]:
        """Get cached OCR result"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT extracted_text, confidence_score, ocr_engine, timestamp
                FROM ocr_cache WHERE file_path = ? AND file_hash = ?
            """, (file_path, file_hash))
            
            result = cursor.fetchone()
            if result:
                return {
                    'extracted_text': result[0],
                    'confidence_score': result[1],
                    'ocr_engine': result[2],
                    'timestamp': result[3]
                }
            return None
    
    def get_supported_formats(self) -> List[Dict[str, Any]]:
        """Get list of supported file formats"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT extension, format_type, processing_method, requires_ocr, description
                FROM supported_formats ORDER BY extension
            """)
            
            return [
                {
                    'extension': row[0],
                    'format_type': row[1],
                    'processing_method': row[2],
                    'requires_ocr': row[3],
                    'description': row[4]
                }
                for row in cursor.fetchall()
            ]
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get file processing statistics"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Total files processed
            cursor.execute("SELECT COUNT(*) FROM file_history")
            total_files = cursor.fetchone()[0]
            
            # Success rate
            cursor.execute("SELECT COUNT(*) FROM file_history WHERE success = 1")
            successful_files = cursor.fetchone()[0]
            
            # OCR usage
            cursor.execute("SELECT COUNT(*) FROM file_history WHERE ocr_used = 1")
            ocr_files = cursor.fetchone()[0]
            
            # File types processed
            cursor.execute("""
                SELECT file_type, COUNT(*) 
                FROM file_history 
                GROUP BY file_type 
                ORDER BY COUNT(*) DESC
            """)
            file_types = dict(cursor.fetchall())
            
            # Average processing time
            cursor.execute("""
                SELECT AVG(processing_time_ms) 
                FROM file_history 
                WHERE processing_time_ms IS NOT NULL
            """)
            avg_processing_time = cursor.fetchone()[0] or 0
            
            return {
                'total_files': total_files,
                'successful_files': successful_files,
                'success_rate': (successful_files / total_files * 100) if total_files > 0 else 0,
                'ocr_files': ocr_files,
                'ocr_usage_rate': (ocr_files / total_files * 100) if total_files > 0 else 0,
                'file_types': file_types,
                'avg_processing_time_ms': round(avg_processing_time, 2)
            }


class EnhancedFileOperations:
    """Enhanced file operations with OCR capabilities"""
    
    def __init__(self):
        self.database = FileOperationsDatabase()
        self.logger = logging.getLogger(__name__)
        
        # Check OCR dependencies
        self.ocr_available = OCR_AVAILABLE
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        # OCR configuration
        self.ocr_config = {
            'lang': 'eng',  # Default language
            'oem': 1,       # OCR Engine Mode
            'psm': 6,       # Page Segmentation Mode
        }
        
        # File size limits (in MB)
        self.max_file_sizes = {
            'text': 50,      # 50MB for text files
            'image': 20,     # 20MB for images
            'pdf': 100,      # 100MB for PDFs
            'document': 50,  # 50MB for documents
        }
    
    def _get_file_hash(self, file_path: str) -> str:
        """Generate hash for file content caching"""
        import hashlib
        try:
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5()
                for chunk in iter(lambda: f.read(4096), b""):
                    file_hash.update(chunk)
                return file_hash.hexdigest()
        except Exception:
            return None
    
    def _check_file_size(self, file_path: str, format_type: str) -> bool:
        """Check if file size is within limits"""
        try:
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # Size in MB
            max_size = self.max_file_sizes.get(format_type, 50)
            return file_size <= max_size
        except Exception:
            return False
    
    def detect_file_type(self, file_path: str) -> Dict[str, Any]:
        """Detect file type and processing method"""
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            # Get mime type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Check supported formats
            formats = self.database.get_supported_formats()
            for fmt in formats:
                if fmt['extension'] == extension:
                    return {
                        'extension': extension,
                        'format_type': fmt['format_type'],
                        'processing_method': fmt['processing_method'],
                        'requires_ocr': fmt['requires_ocr'],
                        'description': fmt['description'],
                        'mime_type': mime_type,
                        'supported': True
                    }
            
            # Unknown file type
            return {
                'extension': extension,
                'format_type': 'unknown',
                'processing_method': 'ocr',  # Try OCR as fallback
                'requires_ocr': True,
                'description': f'Unknown file type: {extension}',
                'mime_type': mime_type,
                'supported': False
            }
            
        except Exception as e:
            return {
                'extension': '',
                'format_type': 'unknown',
                'processing_method': 'error',
                'requires_ocr': False,
                'description': f'Error detecting file type: {e}',
                'mime_type': None,
                'supported': False
            }
    
    def read_text_file(self, file_path: str, encoding: str = 'utf-8') -> Tuple[bool, str, str]:
        """Read plain text file with encoding detection"""
        try:
            # Try specified encoding first
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return True, content, f"Text file read successfully with {encoding} encoding"
            except UnicodeDecodeError:
                # Try other common encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for enc in encodings:
                    try:
                        with open(file_path, 'r', encoding=enc) as f:
                            content = f.read()
                        return True, content, f"Text file read successfully with {enc} encoding"
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, read as binary and decode with errors='replace'
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='replace')
                return True, content, "Text file read with UTF-8 and error replacement"
                
        except Exception as e:
            return False, "", f"Error reading text file: {e}"
    
    def read_pdf_file(self, file_path: str) -> Tuple[bool, str, str]:
        """Read PDF file with text extraction and OCR fallback"""
        if not self.pdf_available:
            return False, "", "PDF processing not available (install pdfplumber)"
        
        try:
            text_content = []
            ocr_used = False
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Try to extract text directly
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        text_content.append(f"--- Page {page_num} ---\n{page_text}\n")
                    else:
                        # No text found, try OCR
                        if self.ocr_available:
                            try:
                                # Convert page to image for OCR
                                page_image = page.to_image(resolution=300)
                                ocr_text = pytesseract.image_to_string(
                                    page_image.original,
                                    config=f'--oem {self.ocr_config["oem"]} --psm {self.ocr_config["psm"]}'
                                )
                                if ocr_text.strip():
                                    text_content.append(f"--- Page {page_num} (OCR) ---\n{ocr_text}\n")
                                    ocr_used = True
                            except Exception as ocr_e:
                                text_content.append(f"--- Page {page_num} ---\n[OCR failed: {ocr_e}]\n")
                        else:
                            text_content.append(f"--- Page {page_num} ---\n[No text extractable, OCR not available]\n")
            
            content = '\n'.join(text_content)
            message = f"PDF processed successfully ({len(pdf.pages)} pages)"
            if ocr_used:
                message += " with OCR fallback"
            
            return True, content, message
            
        except Exception as e:
            return False, "", f"Error reading PDF: {e}"
    
    def read_docx_file(self, file_path: str) -> Tuple[bool, str, str]:
        """Read DOCX file"""
        if not self.docx_available:
            return False, "", "DOCX processing not available (install python-docx)"
        
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content_parts.append(row_text)
            
            content = '\n'.join(content_parts)
            return True, content, f"DOCX file processed successfully ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)"
            
        except Exception as e:
            return False, "", f"Error reading DOCX: {e}"
    
    def read_image_with_ocr(self, file_path: str) -> Tuple[bool, str, str]:
        """Read image file using OCR"""
        if not self.ocr_available:
            return False, "", "OCR not available (install pytesseract and PIL)"
        
        try:
            # Check if we have cached OCR result
            file_hash = self._get_file_hash(file_path)
            if file_hash:
                cached = self.database.get_cached_ocr(file_path, file_hash)
                if cached:
                    return True, cached['extracted_text'], f"OCR result from cache (confidence: {cached.get('confidence_score', 'unknown')})"
            
            # Perform OCR
            image = Image.open(file_path)
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text with confidence
            data = pytesseract.image_to_data(
                image,
                config=f'--oem {self.ocr_config["oem"]} --psm {self.ocr_config["psm"]}',
                output_type=pytesseract.Output.DICT
            )
            
            # Calculate confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Extract text
            text = pytesseract.image_to_string(
                image,
                config=f'--oem {self.ocr_config["oem"]} --psm {self.ocr_config["psm"]}'
            )
            
            # Cache result
            if file_hash and text.strip():
                self.database.cache_ocr_result(file_path, file_hash, text, avg_confidence)
            
            message = f"OCR completed (confidence: {avg_confidence:.1f}%)"
            return True, text, message
            
        except Exception as e:
            return False, "", f"OCR failed: {e}"
    
    def read_structured_data(self, file_path: str) -> Tuple[bool, str, str]:
        """Read structured data files (CSV, JSON, etc.)"""
        try:
            extension = Path(file_path).suffix.lower()
            
            if extension == '.csv':
                if not PANDAS_AVAILABLE:
                    # Fallback to basic CSV reading
                    import csv
                    content_parts = []
                    with open(file_path, 'r', encoding='utf-8') as f:
                        reader = csv.reader(f)
                        for row_num, row in enumerate(reader, 1):
                            content_parts.append('\t'.join(row))
                            if row_num > 1000:  # Limit for large files
                                content_parts.append(f"... (truncated after 1000 rows)")
                                break
                    content = '\n'.join(content_parts)
                    return True, content, f"CSV file read ({row_num} rows)"
                else:
                    # Use pandas for better CSV handling
                    df = pd.read_csv(file_path, nrows=1000)  # Limit rows
                    content = df.to_string()
                    return True, content, f"CSV file read ({len(df)} rows, {len(df.columns)} columns)"
            
            elif extension in ['.json']:
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                content = json.dumps(data, indent=2, ensure_ascii=False)
                return True, content, "JSON file read successfully"
            
            elif extension in ['.xml']:
                # Basic XML reading
                success, content, msg = self.read_text_file(file_path)
                if success:
                    return True, content, "XML file read as text"
                return success, content, msg
            
            elif extension in ['.yaml', '.yml']:
                try:
                    import yaml
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = yaml.safe_load(f)
                    content = yaml.dump(data, default_flow_style=False, allow_unicode=True)
                    return True, content, "YAML file read successfully"
                except ImportError:
                    # Fallback to text reading
                    return self.read_text_file(file_path)
            
            else:
                return False, "", f"Unsupported structured format: {extension}"
                
        except Exception as e:
            return False, "", f"Error reading structured data: {e}"
    
    def read_file(self, file_path: str, force_ocr: bool = False) -> Dict[str, Any]:
        """Main file reading function with automatic format detection"""
        start_time = datetime.now()
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'content': '',
                    'error': f"File not found: {file_path}",
                    'file_type': 'unknown',
                    'processing_method': 'error',
                    'ocr_used': False,
                    'processing_time_ms': 0
                }
            
            # Detect file type
            file_info = self.detect_file_type(file_path)
            file_size = os.path.getsize(file_path)
            
            # Check file size limits
            if not self._check_file_size(file_path, file_info['format_type']):
                max_size = self.max_file_sizes.get(file_info['format_type'], 50)
                return {
                    'success': False,
                    'content': '',
                    'error': f"File too large (max {max_size}MB for {file_info['format_type']} files)",
                    'file_type': file_info['format_type'],
                    'processing_method': 'size_limit',
                    'ocr_used': False,
                    'processing_time_ms': 0
                }
            
            success = False
            content = ""
            message = ""
            ocr_used = False
            processing_method = file_info['processing_method']
            
            # Force OCR if requested
            if force_ocr and self.ocr_available:
                processing_method = 'ocr'
            
            # Process file based on type
            if processing_method == 'text':
                success, content, message = self.read_text_file(file_path)
            
            elif processing_method == 'pdf':
                success, content, message = self.read_pdf_file(file_path)
                ocr_used = 'OCR' in message
            
            elif processing_method == 'docx':
                success, content, message = self.read_docx_file(file_path)
            
            elif processing_method == 'structured':
                success, content, message = self.read_structured_data(file_path)
            
            elif processing_method == 'ocr' or force_ocr:
                success, content, message = self.read_image_with_ocr(file_path)
                ocr_used = True
            
            else:
                # Try text reading as fallback
                success, content, message = self.read_text_file(file_path)
                if not success and self.ocr_available:
                    # Try OCR as last resort
                    success, content, message = self.read_image_with_ocr(file_path)
                    ocr_used = True
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds() * 1000
            
            # Log operation
            self.database.log_file_operation(
                file_path=file_path,
                file_type=file_info['format_type'],
                processing_method=processing_method,
                success=success,
                ocr_used=ocr_used,
                error_message=message if not success else None,
                processing_time_ms=int(processing_time),
                file_size=file_size,
                content_hash=self._get_file_hash(file_path) if success else None
            )
            
            return {
                'success': success,
                'content': content,
                'message': message,
                'file_type': file_info['format_type'],
                'processing_method': processing_method,
                'ocr_used': ocr_used,
                'processing_time_ms': int(processing_time),
                'file_size': file_size,
                'file_info': file_info
            }
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds() * 1000
            
            return {
                'success': False,
                'content': '',
                'error': f"Unexpected error: {e}",
                'file_type': 'unknown',
                'processing_method': 'error',
                'ocr_used': False,
                'processing_time_ms': int(processing_time)
            }
    
    def get_capabilities(self) -> Dict[str, Any]:
        """Get current file processing capabilities"""
        return {
            'ocr_available': self.ocr_available,
            'pdf_available': self.pdf_available,
            'docx_available': self.docx_available,
            'pandas_available': PANDAS_AVAILABLE,
            'supported_formats': self.database.get_supported_formats(),
            'max_file_sizes': self.max_file_sizes,
            'ocr_config': self.ocr_config if self.ocr_available else None
        }
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get file processing statistics"""
        return self.database.get_processing_stats()
    
    def clear_ocr_cache(self) -> bool:
        """Clear OCR cache"""
        try:
            with sqlite3.connect(self.database.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM ocr_cache")
                conn.commit()
            return True
        except Exception:
            return False


# Global instance
enhanced_file_ops = EnhancedFileOperations()


def read_file_enhanced(file_path: str, force_ocr: bool = False) -> Dict[str, Any]:
    """Enhanced file reading with OCR capabilities"""
    return enhanced_file_ops.read_file(file_path, force_ocr)


def get_file_capabilities() -> Dict[str, Any]:
    """Get file processing capabilities"""
    return enhanced_file_ops.get_capabilities()


def get_file_stats() -> Dict[str, Any]:
    """Get file processing statistics"""
    return enhanced_file_ops.get_processing_stats()


def detect_file_type(file_path: str) -> Dict[str, Any]:
    """Detect file type and processing method"""
    return enhanced_file_ops.detect_file_type(file_path)
