import os
import sys
import pytest

# Add the parent folder of `textcleaner_partha` to sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from textcleaner_partha.preprocess import preprocess_file, get_tokens_from_file

# === TEST DATA FILES ===
TXT_FILE = "sample.txt"
DOCX_FILE = "sample.docx"
PDF_FILE = "sample.pdf"
INVALID_FILE = "sample.csv"


@pytest.fixture(scope="module")
def setup_test_files(tmp_path_factory):
    """Fixture to create temporary test files."""
    tmp_dir = tmp_path_factory.mktemp("data")

    # TXT File
    txt_path = tmp_dir / TXT_FILE
    txt_path.write_text("This is a test.\nAnother line here!\nAbbrev: AI.\n")

    # DOCX File
    import docx
    doc = docx.Document()
    doc.add_paragraph("This is a test paragraph in DOCX.")
    doc.add_paragraph("Another paragraph with abbreviations e.g. NLP.")
    doc.save(tmp_dir / DOCX_FILE)

    # PDF File
    from fpdf import FPDF
    pdf_path = tmp_dir / PDF_FILE
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, "This is page 1.\nLine two on page 1.")
    pdf.add_page()
    pdf.multi_cell(0, 10, "This is page 2.\nContains AI and NLP terms.")
    pdf.output(str(pdf_path))

    return str(tmp_dir)


# === TEST CASES ===

def test_preprocess_txt(setup_test_files):
    file_path = os.path.join(setup_test_files, TXT_FILE)
    result = preprocess_file(file_path)
    assert isinstance(result, list)
    assert all(isinstance(line, str) for line in result)
    assert any("test" in line for line in result)


def test_get_tokens_txt(setup_test_files):
    file_path = os.path.join(setup_test_files, TXT_FILE)
    tokens = get_tokens_from_file(file_path)
    assert isinstance(tokens, list)
    assert all(isinstance(token_list, list) for token_list in tokens)


def test_preprocess_docx(setup_test_files):
    file_path = os.path.join(setup_test_files, DOCX_FILE)
    result = preprocess_file(file_path)
    assert isinstance(result, list)
    assert any("paragraph" in line for line in result)


def test_preprocess_pdf_flat(setup_test_files):
    file_path = os.path.join(setup_test_files, PDF_FILE)
    result = preprocess_file(file_path)
    assert isinstance(result, list)
    assert any("page" in line for line in result)


def test_preprocess_pdf_chunked(setup_test_files):
    file_path = os.path.join(setup_test_files, PDF_FILE)
    result = preprocess_file(file_path, pdf_chunk_by_page=True)
    assert isinstance(result, list)
    assert all("page_number" in page for page in result)
    assert all(isinstance(page["content"], list) for page in result)


def test_preprocess_pdf_chunked_merged(setup_test_files):
    file_path = os.path.join(setup_test_files, PDF_FILE)
    result = preprocess_file(file_path, pdf_chunk_by_page=True, merge_pdf_pages=True)
    assert isinstance(result, list)
    assert all(isinstance(line, str) for line in result)


def test_get_tokens_pdf_chunked(setup_test_files):
    file_path = os.path.join(setup_test_files, PDF_FILE)
    tokens = get_tokens_from_file(file_path, pdf_chunk_by_page=True)
    assert isinstance(tokens, list)
    assert all("page_number" in page for page in tokens)


def test_invalid_file_type(setup_test_files):
    file_path = os.path.join(setup_test_files, INVALID_FILE)
    with open(file_path, "w") as f:
        f.write("invalid file type")
    with pytest.raises(ValueError):
        preprocess_file(file_path)


def test_missing_file_error():
    """Test that FileNotFoundError is raised for non-existent files."""
    with pytest.raises(FileNotFoundError):
        preprocess_file("non_existent_file.txt")

    with pytest.raises(FileNotFoundError):
        get_tokens_from_file("non_existent_file.pdf")